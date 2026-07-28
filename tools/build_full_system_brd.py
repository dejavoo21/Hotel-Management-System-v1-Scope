from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "LaFlo_Enterprise_Hotel_Operations_Platform_Full_System_BRD_v1.0.docx"

NAVY = "123B4A"
TEAL = "087F73"
GREEN = "2F8F65"
BLUE = "2B6CB0"
INK = "172B35"
MUTED = "5F6F78"
LIGHT = "F3F7F6"
PALE_TEAL = "E8F5F2"
PALE_BLUE = "EDF4FA"
PALE_AMBER = "FFF7E8"
PALE_RED = "FDEEEE"
BORDER = "D9E3E1"
WHITE = "FFFFFF"
RED = "A33A3A"
AMBER = "9A6A14"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size: float = 10.5, bold: bool = False, color: str = INK, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[object]], widths: Sequence[int], font_size: float = 8.5):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(header))
        set_font(run, size=8.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, LIGHT)
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(str(value))
            set_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_para(doc: Document, text: str, bold_lead: str | None = None, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run, italic=italic)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.167
        run = paragraph.add_run(item)
        set_font(run)


def add_callout(doc: Document, title: str, text: str, tone: str = "teal") -> None:
    fill = {"teal": PALE_TEAL, "blue": PALE_BLUE, "amber": PALE_AMBER, "red": PALE_RED}[tone]
    accent = {"teal": TEAL, "blue": BLUE, "amber": AMBER, "red": RED}[tone]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=accent)
    r2 = p.add_run(f"\n{text}")
    set_font(r2, size=9.5, color=INK)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, NAVY, 16, 8),
        "Heading 2": (13, TEAL, 12, 6),
        "Heading 3": (11.5, BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


MODULES = [
    ("Dashboard", "Executive and operational command view", "GM, operations manager, authorised department users",
     "Arrivals, departures, occupancy, ADR, authorised revenue, room readiness, housekeeping, maintenance, incidents, guest alerts, security, building and integration health, tasks and activity.",
     "Dashboard summaries, tasks, alerts and timeline", "All operational modules, Enterprise Search and Hotel Brain", "Mixed: live API data plus some demo fallbacks requiring explicit labelling"),
    ("Bookings / Reservations", "Manage the complete stay lifecycle", "Front desk, managers",
     "Create, update, cancel, assign room, check in, check out, charges, payments, source, notes and special requests.",
     "Booking, guest, room, charges and payment status", "Guests, Rooms, Calendar, Financials, Messages", "Implemented core"),
    ("Guests", "Maintain the authorised guest record and service history", "Front desk, concierge, managers",
     "Profile, contact details, preferences, stay history, issues, linked records and privacy controls.",
     "Guest profile and journey events", "Bookings, Rooms, Messages, Calls, Reviews, Concierge", "Implemented core"),
    ("Rooms", "Control physical room inventory and readiness", "Front desk, housekeeping, maintenance, managers",
     "Room type, floor, occupancy, cleaning, inspection, out-of-service and blocker states.",
     "Room, room type, floor and linked operational status", "Bookings, Housekeeping, Maintenance, Smart Building", "Implemented core"),
    ("Housekeeping", "Coordinate room turnover and readiness", "Housekeeping managers and staff",
     "Cleaning status, assignment, inspection, priority, delays and blocker escalation.",
     "Room state and housekeeping logs", "Rooms, Bookings, Guests, Maintenance, Calendar", "Partial: no dedicated housekeeping-task entity"),
    ("Inventory", "Control operational stock and replenishment", "Department managers, inventory users",
     "Items, stock levels, thresholds, department ownership, suppliers, purchase orders, movement and low-stock alerts.",
     "Inventory items and purchase orders", "Housekeeping, Maintenance, Reports, Dashboard", "Implemented core; movement/supplier depth requires extension"),
    ("Calendar", "Provide cross-department operational planning", "All authorised operational users",
     "Reservation, arrival, departure, housekeeping, maintenance, incident follow-up, event and reminder views.",
     "Calendar events and linked records", "Bookings, Guests, Rooms, Housekeeping, Maintenance, Incidents, Concierge", "Implemented core; cross-module sync requires validation"),
    ("Financials", "Provide controlled financial visibility", "Finance users, managers, admin",
     "Revenue, invoices, payments, balances, charges, refunds where implemented, ADR and performance by source and room type.",
     "Invoice, charge, payment and computed reports", "Bookings, Guests, Reports, Dashboard, Audit", "Partial: no general ledger/refund/expense domain"),
    ("Reports", "Deliver authorised management and operational reporting", "Managers, finance, security, auditors, admin",
     "Catalogue, generation, preview and export for operational, guest, asset, financial, audit and AI reports.",
     "Computed report results and catalogue metadata", "All modules subject to permission", "Partial: no saved/scheduled report or export-history model"),
    ("Reviews", "Track feedback and service recovery", "Guest experience, concierge, managers",
     "Source, rating, sentiment, response, assignment, links and escalation.",
     "Review records", "Guests, Bookings, Concierge, Operations, Incidents", "Implemented core"),
    ("Concierge", "Coordinate guest service requests", "Concierge, front desk, managers",
     "Request categories, assignment, priority, notes, status and escalation.",
     "Concierge requests", "Guests, Rooms, Bookings, Messages, Calls, Operations", "Implemented core"),
    ("Messages", "Maintain guest and staff conversations", "Front desk, concierge, authorised staff",
     "Threads, linked records, follow-up and escalation; collaboration controls only within an active session.",
     "Conversation, message and ticket records", "Guests, Bookings, Rooms, Concierge, Incidents", "Implemented core"),
    ("Calls", "Support voice interactions and follow-up", "Front desk, concierge, authorised staff",
     "Call session, caller context, notes, links and escalation; call controls only in an active session.",
     "Target call log/session records", "Guests, Bookings, Rooms, Concierge, Incidents", "Partial/placeholder: no persisted CallLog/provider lifecycle"),
    ("User Management", "Govern workforce access", "Admin, system admin, managers for limited reads",
     "Users, account status, access requests, approval, password setup, roles, permissions and disable/reactivate.",
     "User, access request, refresh/reset/OTP records", "All modules, Audit", "Implemented core; role/action policy model needs expansion"),
    ("Operations Center", "Coordinate cross-department issues and workload", "GM, operations manager, department managers",
     "Operational overview, blockers, alerts, tasks, workload, daily status and critical attention items.",
     "Tickets, signals, context, recommendations and task state", "All operational modules", "Partial/hybrid"),
    ("Maintenance Center", "Manage faults and engineering work", "Maintenance managers and staff",
     "Issues, work orders, faults, repairs, preventive work, assets, assignment, priority and room/device blockers.",
     "Maintenance model family", "Rooms, Housekeeping, Smart Building, Incidents, Dashboard", "Implemented core"),
    ("Incident Center", "Control reportable operational and safety incidents", "Operations, security, maintenance, managers",
     "Classification, severity, priority, assignment, investigation, evidence, tasks, escalation and closure.",
     "Incident, task, comment, attachment and SLA records", "All operational and control modules", "Implemented core"),
    ("Security Center", "Provide security and CCTV operational visibility", "Security users, admin",
     "Camera health, access events, security alerts, visitors, restricted-area alerts and incident escalation.",
     "Camera feed, access, door, visitor and alert records", "Integration Manager, Smart Building, Incident Center", "Implemented core; vendor/media adapters partial"),
    ("Smart Building", "Monitor connected hotel assets and environment", "Engineering, security, operations, admin",
     "Device inventory, health, sensors, doors, HVAC, energy, assets, mapping and fault routing.",
     "IoT devices, readings and door status", "Maintenance, Security, Incidents, Integration Manager", "Implemented core; ingestion/adapters partial"),
    ("Integration Manager", "Centralise provider configuration and health", "Integration admin, system admin",
     "Discover, configure, test, map, monitor and disconnect supported providers with masked credentials.",
     "Hardware integration, credential references and runtime provider registry", "Security, Smart Building, Operations, Financials, AI", "Partial: persisted hardware plus environment/in-memory marketplace state"),
    ("Enterprise Search", "Find authorised records across the platform", "All authorised users",
     "Global/advanced search, filters, previews, recent/saved searches and source navigation.",
     "SearchIndex records with access scope", "All indexed modules", "Partial: record filtering implemented; saved search and freshness lifecycle incomplete"),
    ("Hotel Brain / AI", "Provide evidence-based operational intelligence", "GM, managers and authorised users",
     "Natural-language questions, evidence, suggested actions, summaries, sensitive-action confirmation and governance logging.",
     "Search context, AI recommendations and operational records", "All authorised modules", "Partial/rule-provider hybrid; durable query/answer lifecycle incomplete"),
]


ROLE_ROWS = [
    ("General Manager", "Accountable for whole-hotel performance and risk", "Dashboard, Operations, Reports, approved Financials, control centres, Search, Hotel Brain", "Review, prioritise, approve sensitive actions where authorised", "No raw credentials; restricted personal/security evidence only when explicitly granted"),
    ("Operations Manager", "Coordinates daily cross-department operations", "Dashboard, Bookings, Guests, Rooms, Housekeeping, Maintenance, Incidents, Operations, Calendar, authorised reports", "Assign, escalate, close operational work", "Financial, security and admin data require separate permission"),
    ("Front Desk Staff", "Manages bookings, arrival/departure and guest contact", "Bookings, Guests, Rooms, Messages, Calls, Calendar, limited Dashboard", "Create/update bookings and guests, check in/out, record permitted charges", "No financial reporting, security evidence, integrations or user administration"),
    ("Housekeeping Manager", "Owns room-turn and staff workload", "Housekeeping, Rooms, Calendar, Inventory, Maintenance links, limited Dashboard/Reports", "Assign, inspect, prioritise, escalate blockers", "No unrelated guest financial/security data"),
    ("Housekeeping Staff", "Completes assigned cleaning work", "Assigned housekeeping work, room status, limited Calendar/Inventory", "Update own tasks and blockers", "No staffing configuration, financials, security or broad guest data"),
    ("Maintenance Manager", "Owns engineering workload and asset reliability", "Maintenance, Smart Building, Rooms, Inventory, Calendar, Incidents, relevant Reports", "Assign, prioritise, close and escalate", "No raw security/finance data without explicit grant"),
    ("Maintenance Staff", "Completes assigned maintenance work", "Assigned maintenance, relevant asset/room/device context", "Update work, notes and completion evidence", "No configuration or unrelated records"),
    ("Security Manager", "Owns security monitoring and response", "Security, CCTV, access events, relevant Smart Building, Incidents, Reports", "Investigate, escalate, close, manage authorised camera mappings", "No financials or integration secrets"),
    ("Security Staff", "Monitors and handles assigned security events", "Security Center, assigned incidents and approved CCTV", "Acknowledge, investigate, add evidence", "No provider credentials or unrelated guest/finance records"),
    ("Finance User", "Manages authorised financial records", "Financials, finance Reports, linked booking/guest summaries", "Record payments, invoices, authorised refunds/exports", "No security, audit or full guest profile without separate grant"),
    ("Concierge Staff", "Coordinates guest requests and service recovery", "Concierge, Guests, Messages, Calls, Reviews, linked Bookings/Rooms", "Assign, update, escalate and communicate", "No financial/security/admin access"),
    ("IT / Integration Administrator", "Configures platform and provider connectivity", "Settings, Integration Manager, device/provider health, limited Audit", "Configure/test providers, manage mappings and references", "Secrets remain masked; no operational/financial data unless needed and granted"),
    ("System Administrator", "Manages users, platform configuration and access", "Users, Settings, Integrations, Audit and all modules by controlled superuser policy", "Provision, disable, assign permissions, administer platform", "Use least privilege; privileged actions fully audited"),
    ("Auditor", "Reviews evidence and control operation", "Read-only Audit, authorised Reports, access history", "Search, review and export approved evidence", "No mutation; no secrets; data minimisation applies"),
    ("AI Governance Reviewer", "Reviews AI recommendations and sensitive actions", "AI governance, authorised source records and audit", "Approve/reject recommendations and record rationale", "Cannot access source records outside own permissions"),
    ("Department Manager", "Owns a department's workload and service levels", "Department modules, Dashboard aggregates, Calendar, relevant Reports", "Assign, approve and report within department", "No cross-department sensitive data without grant"),
]


PROCESS_ROWS = [
    ("Reservation to check-in", "Booking created; guest linked; room assigned; readiness and blockers validated; check-in completed", "Bookings, Guests, Rooms, Housekeeping, Maintenance, Dashboard, Audit"),
    ("Check-out to room turnover", "Check-out marks room dirty; cleaning work assigned; inspection completes; room returns ready", "Bookings, Rooms, Housekeeping, Calendar, Dashboard"),
    ("Housekeeping task management", "Work created/derived; assigned; progressed; inspected; delayed/blocker conditions escalated", "Housekeeping, Rooms, Maintenance, Notifications"),
    ("Guest request handling", "Request captured; contextualised; assigned; fulfilled or escalated; guest updated", "Concierge, Guests, Messages, Calls, Operations"),
    ("Guest complaint handling", "Complaint captured; priority assessed; ticket/incident created; owner resolves; closure communicated", "Reviews, Messages, Calls, Concierge, Operations, Incidents"),
    ("Maintenance issue handling", "Fault captured manually/device; triaged; assigned; completed; room/device state restored", "Maintenance, Rooms, Smart Building, Incidents"),
    ("Incident reporting and escalation", "Incident classified; severity/priority set; evidence and tasks maintained; closure controlled", "Incident Center and linked modules"),
    ("Security alert handling", "Alert acknowledged; investigated; evidence linked; incident escalated/closed", "Security, CCTV, Smart Building, Incidents"),
    ("CCTV health handling", "Offline/degraded feed detected; security notified; integration checked; fault escalated", "Security, Integration Manager, Maintenance"),
    ("Smart building fault handling", "Device event validated; mapped context resolved; maintenance/security/incident route selected", "Smart Building, Maintenance, Security, Incidents"),
    ("Inventory low-stock handling", "Threshold breach identified; department notified; reorder suggested/task created; PO processed", "Inventory, Notifications, Operations, Reports"),
    ("Review handling", "Review ingested; sentiment and ownership set; response tracked; serious issue escalated", "Reviews, Guests, Operations, Incidents"),
    ("Payment/financial handling", "Charge/invoice/payment recorded; status updated; exception resolved; audit written", "Financials, Bookings, Guests, Audit"),
    ("Integration setup and monitoring", "Provider chosen; credentials referenced; test run; resources mapped; health monitored", "Integration Manager and consuming modules"),
    ("Enterprise Search investigation", "Query submitted; record access filtered; preview opened; source record accessed; audit written", "Enterprise Search and indexed modules"),
    ("Hotel Brain operational briefing", "Authorised context assembled; answer/evidence produced; actions proposed; confirmation governed", "Hotel Brain, AI Context, Search, Audit"),
    ("Daily GM briefing", "Daily context consolidated; priority risks and attention items summarised for authorised GM", "Dashboard, Hotel Brain, AI Context, all control modules"),
]


JOURNEYS = [
    ("Journey 1: Reservation to Check-In",
     ["Create the booking and capture source, dates, room type, notes and requests.",
      "Create or select the guest profile and link it to the booking.",
      "Assign a suitable room and validate occupancy conflicts.",
      "Check room cleaning, inspection and out-of-service status.",
      "Check unresolved maintenance blockers and required guest requests.",
      "Complete check-in using an authorised front-desk action.",
      "Update booking, room, dashboard, calendar and guest journey state.",
      "Create audit entries and relevant notifications."],
     "Bookings, Guests, Rooms, Housekeeping, Maintenance, Calendar, Dashboard, Audit"),
    ("Journey 2: Check-Out to Room Turnover",
     ["Complete checkout and final permitted financial status checks.",
      "Change room to dirty/awaiting cleaning and release booking occupancy.",
      "Create or suggest a housekeeping task with priority based on next arrival.",
      "Assign staff and expose work in the housekeeping schedule.",
      "Complete cleaning and record any maintenance blocker.",
      "Perform inspection and resolve failed inspection if needed.",
      "Mark room ready and refresh availability, dashboard and calendar.",
      "Audit all status transitions."],
     "Bookings, Financials, Rooms, Housekeeping, Maintenance, Calendar, Dashboard"),
    ("Journey 3: Guest Complaint to Resolution",
     ["Capture complaint through message, call, review or concierge request.",
      "Link the guest, booking, room and communication evidence.",
      "Classify urgency and alert Operations Center.",
      "Create a task or incident according to severity and policy.",
      "Assign an owner and track response and SLA.",
      "Resolve the issue and communicate the outcome.",
      "Update guest journey/service history.",
      "Surface unresolved or recurring risk in Dashboard and Hotel Brain."],
     "Messages, Calls, Reviews, Concierge, Guests, Operations, Incidents, Hotel Brain"),
    ("Journey 4: Maintenance Fault to Closure",
     ["Capture fault manually or receive a mapped device event.",
      "Resolve room/device/area context and determine guest impact.",
      "Create a maintenance record and assign priority.",
      "Assign technician and apply any room/device blocker.",
      "Record diagnosis, work and parts used.",
      "Complete and verify work; restore room/device health.",
      "Close work and any linked incident.",
      "Write audit, dashboard and notification updates."],
     "Maintenance, Smart Building, Rooms, Inventory, Incidents, Dashboard"),
    ("Journey 5: Security Event to Incident Closure",
     ["Receive door, access, CCTV or security alert.",
      "Present the event in Security Center to an authorised user.",
      "Acknowledge and assess severity.",
      "Create and link an incident when policy requires.",
      "Add investigation notes and authorised evidence.",
      "Assign follow-up actions and track resolution.",
      "Close the alert/incident with required rationale.",
      "Record audit and management notifications."],
     "Security, CCTV, Smart Building, Incidents, Notifications, Audit"),
    ("Journey 6: CCTV Integration to Security Center",
     ["Open Settings > Integrations > CCTV with configure permission.",
      "Select local, ONVIF, manual, NVR or supported provider method.",
      "Provide credential references and connection details securely.",
      "Test connection and display actionable errors.",
      "Discover/import cameras or channels and label simulation when used.",
      "Map camera feeds to hotel areas.",
      "Expose authorised camera health in Security Center.",
      "Monitor health and audit connection/mapping actions."],
     "Settings, Integration Manager, CCTV, Security Center, Audit"),
    ("Journey 7: Smart Building Device Setup to Alert Routing",
     ["Open the Smart Building integration category.",
      "Select provider/category and securely configure a connection.",
      "Test the connection and discover/import devices.",
      "Map devices to rooms, floors or areas.",
      "Receive and validate a device event.",
      "Apply threshold/routing policy.",
      "Create Maintenance, Security or Incident work as appropriate.",
      "Update device health, notifications and audit."],
     "Integration Manager, Smart Building, Maintenance, Security, Incidents"),
    ("Journey 8: Enterprise Search Investigation",
     ["Enter a search term or structured filter.",
      "Search only records within the user's module and record scope.",
      "Group results by category and show safe previews.",
      "Open the source record through its protected route.",
      "Record search and restricted-attempt audit events where policy requires."],
     "Enterprise Search, RBAC, indexed modules, Audit"),
    ("Journey 9: Hotel Brain Operational Question",
     ["Submit a natural-language operational question.",
      "Assemble only authorised and current context.",
      "Return a bounded answer with supporting records/evidence.",
      "Present uncertainty or insufficient-data state when appropriate.",
      "Suggest actions without executing high-impact changes.",
      "Require explicit confirmation and permission for sensitive actions.",
      "Log the question, evidence references, answer and decision."],
     "Hotel Brain, AI Context, Enterprise Search, Governance, Audit"),
    ("Journey 10: Daily GM Briefing",
     ["Aggregate authorised daily operational context.",
      "Identify not-ready rooms, open incidents, maintenance blockers, guest alerts, device issues and integration health.",
      "Prioritise by impact and urgency.",
      "Show evidence links and suggested priorities.",
      "Protect department-sensitive and financial detail.",
      "Record generation and user access where policy requires."],
     "Daily GM Briefing, Hotel Brain, Dashboard, all control modules"),
]


FR_TEMPLATES = {
    "DASH": ("Dashboard", ["display today's operational KPIs", "show room-readiness exceptions", "show authorised cross-department alerts", "hide financial values without financials permission", "link each attention item to an authorised source record", "label demo or simulated widget data", "show loading, empty and error states"]),
    "BOOK": ("Bookings", ["create and validate a booking", "link a guest and room", "update or cancel within role policy", "complete check-in and check-out", "record notes and special requests", "surface payment status without exposing restricted detail", "publish lifecycle events and audit entries"]),
    "GUEST": ("Guests", ["create and update a guest profile", "maintain contact details and preferences", "show stay history and linked records", "restrict personal data by permission", "record complaints and journey events", "support safe search and source navigation"]),
    "ROOM": ("Rooms", ["maintain room, type and floor data", "track occupancy and readiness states", "apply out-of-service and maintenance blockers", "link active guest and booking", "link housekeeping and smart-device context", "audit material status changes"]),
    "HK": ("Housekeeping", ["create or derive cleaning work", "assign work to staff", "track cleaning and inspection status", "prioritise rooms by arrival need", "escalate delays and blockers", "update room readiness and calendar", "notify responsible users"]),
    "INV": ("Inventory", ["maintain item and stock records", "set reorder thresholds", "flag low and out-of-stock items", "record department and supplier context", "create reorder suggestion or purchase order", "export authorised inventory data", "notify on critical stock conditions"]),
    "CAL": ("Calendar", ["show operational events in month, week, day and agenda views", "filter by type, department, room and status", "link events to source records", "create and update authorised events", "reflect booking and task status changes", "provide empty/loading/error states"]),
    "FIN": ("Financials", ["show revenue and ADR only to authorised users", "manage invoices, charges and payments", "track balances and payment status", "support refunds only where implemented and approved", "link records to guest and booking", "export authorised financial data", "audit financial changes and exports"]),
    "REP": ("Reports", ["show a permission-filtered report catalogue", "generate reports from authorised data", "preview and export supported formats", "restrict finance, audit and security reports", "label prepared/demo results", "record generation and export audit", "support saved/scheduled reports when implemented"]),
    "REV": ("Reviews", ["capture review source, rating and text", "track sentiment and response status", "assign an owner", "link guest and booking where authorised", "escalate serious feedback", "track response and closure"]),
    "CON": ("Concierge", ["create and classify a guest request", "set priority, owner and status", "link guest, room and booking", "maintain notes and communications", "escalate to operations, maintenance or incident", "notify owner and requester"]),
    "MSG": ("Messages", ["maintain conversation threads", "link authorised guest/booking/room context", "create follow-up tasks or requests", "escalate serious content", "show session toolbar only in active collaboration", "audit material actions"]),
    "CALL": ("Calls", ["create a persisted call session/log", "resolve caller and linked hotel records", "record notes and follow-up", "escalate to concierge or incident", "show call controls only in active session", "handle provider unavailable and failed-call states"]),
    "USER": ("User Management", ["create and update users", "approve or reject access requests", "assign roles and module/action permissions", "disable and reactivate accounts", "track password-setup status", "audit privileged changes", "restrict administration to authorised roles"]),
    "OPS": ("Operations Center", ["aggregate cross-department blockers and alerts", "show workload and ownership", "create and assign operational tasks", "prioritise daily attention items", "link to source modules", "protect restricted data"]),
    "MAINT": ("Maintenance Center", ["create and classify maintenance work", "assign priority and technician", "apply room/device blockers", "track work, parts and evidence", "manage preventive work and assets", "close and audit work"]),
    "INC": ("Incident Center", ["create and classify incidents", "set severity, priority, owner and SLA", "link affected records and evidence", "maintain investigation notes and tasks", "control escalation and closure", "audit every material transition"]),
    "SEC": ("Security Center", ["show authorised alerts and access events", "show camera health without exposing credentials", "acknowledge and investigate alerts", "escalate to incidents", "restrict security evidence", "audit viewing and changes where required"]),
    "SB": ("Smart Building", ["maintain device inventory and mapping", "show current device health", "ingest and validate sensor readings", "route faults to responsible modules", "show energy/HVAC/door states", "label simulated devices and readings"]),
    "INT": ("Integration Manager", ["catalogue supported providers", "store credential references securely", "test connections and return actionable errors", "discover/import/map resources", "show configured, tested, connected and healthy as separate states", "audit configuration actions", "label mocked/planned providers"]),
    "SEARCH": ("Enterprise Search", ["index authorised record metadata", "filter every result by access scope", "support category and structured filters", "show safe previews and source links", "handle no results and unavailable index", "audit restricted attempts and selected searches"]),
    "AI": ("Hotel Brain and AI", ["assemble permission-filtered context", "answer with supporting evidence", "state uncertainty and insufficient data", "suggest bounded actions", "require human confirmation for sensitive actions", "record governance decisions", "prevent disclosure through prompts or summaries"]),
}


NFR_ROWS = [
    ("NFR-001", "Performance", "Primary authenticated pages should reach usable state within 3 seconds at the agreed test load, excluding unavailable third-party providers.", "Must"),
    ("NFR-002", "API performance", "95% of core non-integration API reads should complete within 1 second in the agreed test environment.", "Should"),
    ("NFR-003", "Availability", "Production availability target and maintenance windows shall be agreed; health endpoints shall distinguish platform and dependency health.", "Must"),
    ("NFR-004", "Reliability", "Retryable operations shall be idempotent or protected against duplicate records.", "Must"),
    ("NFR-005", "Scalability", "The platform shall support horizontal API scaling without relying on process-local business state.", "Should"),
    ("NFR-006", "Security", "All protected endpoints shall authenticate and authorise server-side; frontend hiding is not a security boundary.", "Must"),
    ("NFR-007", "Privacy", "Personal, financial, security and credential data shall be minimised and restricted by purpose.", "Must"),
    ("NFR-008", "Accessibility", "Core workflows shall target WCAG 2.2 AA for keyboard use, focus, labels, contrast and error identification.", "Must"),
    ("NFR-009", "Usability", "Common operational actions shall use consistent controls, status language and confirmation patterns.", "Should"),
    ("NFR-010", "Maintainability", "Module, permission, route and provider registries shall have a single documented source of truth.", "Must"),
    ("NFR-011", "Auditability", "Material and privileged actions shall create immutable server-side audit records with actor, target, action and timestamp.", "Must"),
    ("NFR-012", "Data integrity", "State transitions shall validate current state, hotel scope and related-record constraints.", "Must"),
    ("NFR-013", "Error handling", "Known states shall return specific user-safe messages and stable error codes.", "Must"),
    ("NFR-014", "Responsive design", "Core pages shall remain usable at desktop, smaller laptop and tablet widths; mobile support shall match agreed scope.", "Should"),
    ("NFR-015", "Browser compatibility", "Current stable Chrome, Edge and Safari shall be supported; Firefox support shall be agreed.", "Should"),
    ("NFR-016", "Integration resilience", "Provider failures shall not block unrelated modules; timeouts, retry and degraded states shall be visible.", "Must"),
    ("NFR-017", "Credential security", "Secrets shall be encrypted or held by a secure external store and never returned in plaintext.", "Must"),
    ("NFR-018", "Observability", "Services shall emit correlation identifiers, structured errors, health metrics and actionable operational logs.", "Should"),
    ("NFR-019", "AI governance", "AI outputs shall preserve permissions, evidence, uncertainty and human approval for high-impact actions.", "Must"),
    ("NFR-020", "Recovery", "Backup, restore, retention and recovery objectives shall be documented and tested before production sign-off.", "Must"),
]


DATA_ROWS = [
    ("Guest", "Person receiving hotel services", "ID, name, contact, preferences, consent/privacy attributes", "Guests", "Booking, Message, Review, Concierge", "Personal data; least privilege"),
    ("Booking / Reservation", "Planned or active stay", "Reference, guest, room/type, dates, status, source, totals, notes", "Bookings", "Guest, Room, Charge, Payment", "Hotel scope; role/action controls"),
    ("Room", "Physical sellable room", "Number, type, floor, occupancy, cleaning, inspection, service state", "Rooms", "Booking, housekeeping, maintenance, device", "Operational roles"),
    ("Housekeeping Task", "Assigned room-turn work", "Target room, task type, priority, assignee, due time, status, blockers", "Housekeeping", "Room, Booking, User, Maintenance", "Department scope; target model required"),
    ("Maintenance Task", "Engineering work record", "Asset/room, issue, priority, assignee, status, evidence, dates", "Maintenance", "Room, Device, Incident, Inventory", "Maintenance and management"),
    ("Incident", "Controlled reportable event", "Type, severity, priority, status, owner, SLA, evidence, links", "Incident Center", "Tasks, comments, attachments, source records", "Restricted by incident/security scope"),
    ("Inventory Item", "Tracked supply or spare", "Name, category, department, quantity, threshold, unit, cost, supplier", "Inventory", "Purchase order, department task", "Cost restricted where applicable"),
    ("Calendar Event", "Operational schedule entry", "Title, type, times, owner, status, room/area, source link", "Calendar", "Booking, room, task, incident", "Source-record permission"),
    ("Financial Record", "Charge, invoice, payment, refund or adjustment", "ID, amount, currency, method, status, dates, booking/guest", "Financials", "Booking, Guest, Audit", "Financial permission; partial current model"),
    ("Report", "Generated or saved reporting output", "Type, filters, owner, format, created date, access scope", "Reports", "All source modules", "Source and export permissions; persistence gap"),
    ("Review", "Guest feedback record", "Source, rating, text, sentiment, status, owner", "Reviews", "Guest, Booking, Incident", "Guest experience scope"),
    ("Concierge Request", "Guest service request", "Category, priority, status, owner, notes, guest/room/booking", "Concierge", "Guest, Booking, Room, Ticket", "Service team scope"),
    ("Message", "Communication in a conversation", "Thread, sender, body, time, channel, linked context", "Messages", "Conversation, Guest, Booking", "Participant/module scope"),
    ("Call Log", "Persisted call lifecycle and outcome", "Session ID, parties, direction, times, status, notes, recording ref", "Calls", "Guest, Booking, Concierge, Incident", "Highly restricted; current model missing"),
    ("User", "Workforce account", "ID, identity, role, department, status, hotel, password state", "User Management", "Permission, access request, audit", "Admin/manager controls"),
    ("Role / Permission", "Access policy assignment", "Role, module grants, action grants, record scope", "User Management", "User, route/action policies", "Server-authoritative; normalized model gap"),
    ("Access Request", "Request for platform access", "Applicant, requested role/scope, status, approver, dates", "Authentication/Admin", "User, audit, reply", "Admin only after submission"),
    ("Integration", "Provider connection configuration", "Provider, category, state, hotel, capability, health", "Integration Manager", "Credential reference, resources", "Configure permission; generic model partial"),
    ("Credential Reference", "Non-plaintext reference to provider secret", "Reference, masked label, secret store key, rotation metadata", "Integration Manager", "Integration", "Never return plaintext"),
    ("CCTV Camera / NVR Channel", "Video source metadata", "ID, provider, host reference, channel, area, health, stream ref", "Security/CCTV", "Integration, incident", "Security only; current CameraFeed/partial adapters"),
    ("Smart Building Device", "Connected physical/virtual asset", "ID, category, provider, mapping, status, last seen", "Smart Building", "Readings, room/floor, integration", "Engineering/security scope"),
    ("Sensor / Energy / HVAC Reading", "Time-series device measurement", "Device, metric, value, unit, timestamp, quality", "Smart Building", "Device, alert", "Scoped; retention required"),
    ("Door / Security Event", "Access or security observation", "Device/door, event type, actor/ref, result, time, area", "Security", "Alert, incident", "Restricted security evidence"),
    ("Audit Log", "Immutable record of material action", "Actor, action, target, before/after ref, time, result, correlation", "Platform Core", "All modules", "Auditor/admin read; immutable"),
    ("Notification", "Routed user alert", "Type, priority, recipient, channel, payload ref, status, escalation", "Notification Engine", "Event, task, user", "Recipient scoped"),
    ("Task", "Cross-module unit of work", "Type, source, owner, priority, due, status, links", "Task Engine", "Events, users, source records", "Department/record scope"),
    ("Search Index Record", "Search-safe representation of a source record", "Source type/ID, hotel, text, metadata, access scope, indexed at", "Enterprise Search", "Source record", "Permission-filtered; no secret payload"),
    ("AI Recommendation", "Governed suggested action", "Context refs, recommendation, risk, status, reviewer, rationale", "AI Governance", "Source records, audit, task", "Reviewer plus source permissions"),
    ("Hotel Brain Query / Answer", "Auditable AI interaction", "User, question, context/evidence refs, answer, uncertainty, actions, time", "Hotel Brain", "Search records, recommendation, audit", "Current durable model missing"),
]


INTEGRATIONS = [
    ("CCTV / ONVIF", "Camera discovery and health", "Axis, Hikvision, Dahua, generic ONVIF", "ONVIF/manual network configuration", "Camera metadata, health, events", "Discovery, test, offline", "Encrypted references; security permission; network isolation", "Partial; simulation available"),
    ("NVR", "Import channels from recording systems", "Hikvision, Dahua, generic NVR", "Vendor API/ONVIF/manual", "Channel metadata and health", "Test, import, channel state", "No raw credentials/URLs; secure gateway", "Mocked/reference-only channels"),
    ("Local Camera", "Browser-local diagnostic preview", "USB webcam", "Browser MediaDevices", "Local video only", "User-initiated test", "Explicit camera permission; no silent persistence", "Implemented local test"),
    ("Smart Locks", "Door state and access control integration", "TTLock, SALTO, generic", "REST/webhook/vendor gateway", "Locks, access events, status", "Access event, offline, battery", "Signed callbacks; least privilege", "Partial/future adapters"),
    ("Sensors", "Environmental and occupancy telemetry", "Generic IoT, vendor APIs", "REST, webhook, MQTT", "Readings and device health", "Threshold and offline events", "Device identity and signed ingestion", "Partial"),
    ("HVAC", "Monitor climate equipment", "BMS/HVAC providers", "BACnet, Modbus, MQTT, REST", "Set points, readings, faults", "Fault, threshold, offline", "Network segmentation; controlled commands", "Planned/partial model"),
    ("Energy Meters", "Track consumption and anomalies", "Meters/BMS", "Modbus, BACnet, MQTT, REST", "Consumption, demand, health", "Reading, anomaly, offline", "Device identity; retention controls", "Planned/partial model"),
    ("Weather", "Provide operational weather context", "OpenWeather", "REST API", "Current/forecast conditions", "Scheduled refresh, severe condition", "API key protected; graceful degradation", "Environment-configured"),
    ("Payments", "Process and reconcile payments", "Stripe", "Provider API/webhooks", "Payment intent/status/reference", "Payment success/failure/refund", "PCI scope minimisation; signed webhooks", "Partial/environment-configured"),
    ("Booking Channels", "Exchange reservations and inventory", "Booking.com, Expedia", "Certified channel API/webhooks", "Bookings, availability, rates", "Create/update/cancel", "Signed requests, reconciliation", "Coming soon"),
    ("Microsoft 365", "Calendar/email collaboration", "Microsoft 365", "OAuth/Graph API", "Events, mail/task references", "Authorised sync", "OAuth scopes, token rotation", "Planned/partial registry"),
    ("Google Workspace", "Calendar/email collaboration", "Google Workspace", "OAuth APIs", "Events and message references", "Authorised sync", "OAuth scopes, token rotation", "Planned/partial registry"),
    ("OpenAI / AI Provider", "Generate governed summaries/answers", "OpenAI or approved provider", "HTTPS API", "Authorised prompts/context and outputs", "Question, briefing, recommendation", "Data minimisation, provider policy, no secrets", "Optional/environment-configured"),
    ("Generic REST API", "Future provider connectivity", "Approved vendors", "HTTPS REST", "Provider-specific records", "Polling/webhook/action", "Authentication, allow-list, validation", "Planned framework"),
    ("Webhooks", "Receive near-real-time provider events", "Approved providers", "HTTPS callback", "Signed event payloads", "Provider events", "Signature, replay protection, rate limits", "Partial"),
    ("MQTT", "Receive building telemetry", "IoT gateways", "MQTT/TLS", "Device readings and state", "Topic messages", "mTLS/credentials, topic ACL", "Registry/future adapter"),
    ("BACnet", "Connect building automation", "BMS providers", "BACnet gateway", "HVAC/energy/device points", "Point changes", "Gateway isolation, read-only default", "Planned"),
    ("Modbus", "Connect industrial meters/controllers", "Meters/controllers", "Modbus gateway", "Registers mapped to metrics", "Polling/threshold", "Network isolation, safe command policy", "Planned"),
]


REPORT_ROWS = [
    ("Executive dashboard", "Daily high-level KPIs, readiness, risk and attention", "GM/authorised management", "Live aggregates; demo clearly labelled", "Drill-through to authorised source"),
    ("Operations dashboard", "Workload, blockers, escalations and tasks", "Operations/department managers", "Live operational data", "Department and date filters"),
    ("Room readiness", "Ready/not-ready/dirty/blocked/inspection states", "Front desk, housekeeping, managers", "Rooms and housekeeping", "Room and blocker drill-through"),
    ("Reservations", "Volume, status, source, arrivals/departures", "Front desk/managers", "Bookings", "CSV/PDF where supported"),
    ("Housekeeping", "Workload, delay, inspection and productivity", "Housekeeping/managers", "Housekeeping/rooms", "Staff privacy controls"),
    ("Maintenance", "Open, overdue, priority, asset and room blockers", "Maintenance/managers", "Maintenance", "Evidence links"),
    ("Incident", "Volume, severity, SLA, ownership and closure", "Incident-authorised users", "Incidents", "Restricted evidence"),
    ("Security", "Alerts, access events and trends", "Security/admin", "Security Center", "Restricted"),
    ("Smart Building", "Device health, telemetry exceptions and building score", "Engineering/management", "Smart Building", "Simulation labels"),
    ("CCTV health", "Online/offline/degraded camera/channel status", "Security/integration admin", "CCTV/Integration Manager", "No stream credentials"),
    ("Integration health", "Configured/tested/connected/healthy providers", "Integration admin/management", "Integration Manager", "State definitions explicit"),
    ("Guest experience", "Reviews, requests, complaints and resolution", "Guest experience/managers", "Reviews/Concierge/Operations", "Personal data minimised"),
    ("Financial", "Revenue, ADR, payments, balances and source/type performance", "Finance/authorised managers", "Financials", "Financial permission required"),
    ("Inventory", "Stock, thresholds, valuation where authorised and POs", "Inventory/department managers", "Inventory", "Cost permission"),
    ("Audit", "Privileged and material activity evidence", "Auditor/admin", "ActivityLog", "Read-only and export-controlled"),
    ("AI recommendation", "Recommendation state, evidence, reviewer and outcome", "AI governance/authorised managers", "AI Governance", "Source permissions preserved"),
]


AI_ROWS = [
    ("AI Context Engine", "Assemble current, permission-filtered operational context", "Authorised platform records and events", "AI services/system", "Context package with evidence references", "No independent decisions; freshness limits", "Context access and generation logged"),
    ("Daily GM Briefing", "Summarise daily hotel priorities", "Rooms, bookings, incidents, maintenance, guest alerts, devices, integrations", "GM/authorised management", "Briefing, risks and suggested priorities", "Only available data; financial detail permissioned", "Generation/access logged"),
    ("Department Intelligence", "Summarise department workload and risk", "Department records and SLAs", "Department managers", "Department-specific insight", "No cross-department leakage", "Source/evidence and access logged"),
    ("AI Recommendation Governance", "Control recommendation approval and execution", "Recommendations and source records", "Governance reviewer/authorised manager", "Approve, reject, rationale, bounded execution", "Human approval required for sensitive actions", "Full decision audit"),
    ("AI Copilot", "Assist users in current workflow", "Current module context", "Authorised operational users", "Guidance, draft or suggested next action", "Cannot bypass validation or permission", "Material suggestions/actions logged"),
    ("Operations Concierge", "Route operational questions and tasks", "Operations context, tickets and assignments", "Operations users", "Answer, routing suggestion or task draft", "Human confirmation for creation/escalation", "Prompt/output/action audit"),
    ("Hotel Brain", "Answer cross-platform operational questions with evidence", "Enterprise Search index and authorised records", "GM/managers/authorised users", "Answer, evidence, uncertainty and actions", "No hidden data; insufficient-data state required", "Question, evidence, answer and action logged"),
]


NOTIFICATIONS = [
    ("Booking update", "Booking create/change/cancel", "Front desk and relevant owner", "In-app; email/SMS where configured", "Normal", "Escalate only for arrival-impacting failure"),
    ("Check-in/check-out", "Lifecycle completion or blocking error", "Front desk/operations", "In-app", "Normal/High", "High if room or payment blocker"),
    ("Room readiness blocker", "Room not ready near arrival", "Housekeeping manager, front desk, operations", "In-app/push", "High", "Escalate to manager by SLA"),
    ("Housekeeping delay", "Task overdue or failed inspection", "Assignee and housekeeping manager", "In-app", "Medium/High", "Escalate by arrival impact"),
    ("Maintenance fault", "Fault created/device fault", "Maintenance owner/manager", "In-app/push", "Based on priority", "Critical to operations/incident"),
    ("Incident escalation", "Severity/SLA/escalation change", "Incident owner and management/security as scoped", "In-app/email/push", "High/Critical", "Escalation matrix"),
    ("Security alert", "Access/door/security rule event", "Security team", "In-app/push", "High/Critical", "Create incident by policy"),
    ("CCTV offline", "Feed/channel health degraded", "Security and integration admin", "In-app", "Medium/High", "Escalate by coverage criticality"),
    ("Device offline", "Missed heartbeat", "Maintenance/integration admin", "In-app", "Medium", "Escalate by asset class"),
    ("Low battery", "Battery threshold breach", "Maintenance", "In-app", "Low/Medium", "Escalate before failure threshold"),
    ("Water leak", "Leak sensor event", "Maintenance, operations, security if required", "Push/SMS/in-app", "Critical", "Immediate incident/room block"),
    ("Smoke/fire", "Verified life-safety event", "Security/operations/emergency workflow", "Configured critical channels", "Critical", "Follow approved life-safety process; no AI autonomy"),
    ("Inventory low stock", "Quantity at/below threshold", "Department owner/inventory manager", "In-app/email digest", "Medium", "Escalate if critical item/out of stock"),
    ("Integration failure", "Connection/test/health failure", "Integration admin and consuming owner", "In-app/email", "Medium/High", "Escalate if critical service unavailable"),
    ("Guest complaint", "Complaint or service-recovery event", "Concierge/operations owner", "In-app", "Based on impact", "Incident for serious cases"),
    ("Negative review", "Rating/sentiment threshold", "Guest experience/operations", "In-app/digest", "Medium/High", "Escalate unresolved high impact"),
    ("Hotel Brain attention item", "Governed rule identifies material risk", "Authorised manager", "In-app briefing", "Advisory", "Never replaces source alert"),
]


EVENT_ROWS = [
    ("Booking", "booking.created, booking.updated, booking.cancelled, booking.checked_in, booking.checked_out", "Bookings", "Dashboard, rooms, calendar, notifications, audit"),
    ("Guest", "guest.created, guest.updated, guest.issue_logged", "Guests", "Search, guest journey, audit"),
    ("Room", "room.status_changed, room.blocked, room.ready", "Rooms", "Housekeeping, bookings, dashboard"),
    ("Housekeeping", "housekeeping.assigned, started, delayed, inspected, completed", "Housekeeping", "Rooms, calendar, notifications"),
    ("Maintenance", "maintenance.created, assigned, blocked_room, completed, overdue", "Maintenance", "Rooms, incidents, notifications"),
    ("Incident", "incident.created, escalated, updated, closed, sla_breached", "Incident Center", "Operations, security, notifications"),
    ("Security", "security.alert_created, access_denied, door_forced, alert_closed", "Security", "Incidents, notifications, audit"),
    ("Smart Building", "device.discovered, device.offline, reading.threshold_breached, device.faulted", "Smart Building", "Maintenance, security, incidents"),
    ("CCTV", "camera.discovered, camera.offline, stream_tested, camera.mapped", "CCTV/Integrations", "Security, maintenance, audit"),
    ("Integration", "integration.configured, tested, connected, failed, disconnected", "Integration Manager", "Health, notifications, audit"),
    ("Inventory", "inventory.adjusted, low_stock, out_of_stock, purchase_order_created", "Inventory", "Operations, notifications, reports"),
    ("Review", "review.received, assigned, responded, escalated", "Reviews", "Operations, incidents, notifications"),
    ("Message / Call", "message.received, conversation.escalated, call.started, call.ended, followup.created", "Communications", "Concierge, incidents, audit"),
    ("Notification", "notification.created, delivered, failed, acknowledged, escalated", "Notification Engine", "Users, audit"),
    ("Task", "task.created, assigned, started, completed, overdue, cancelled", "Task Engine", "Departments, calendar, notifications"),
    ("Enterprise Search", "search.executed, search.restricted, index.updated, index.failed", "Enterprise Search", "Audit, operations"),
    ("Hotel Brain", "ai.question_asked, answer_generated, recommendation_created, approved, rejected, executed", "AI", "Governance, tasks, audit"),
    ("Audit", "audit.recorded, audit.write_failed", "Audit Engine", "Operations/security monitoring"),
]


AC_ROWS = []
for prefix, (module, capabilities) in FR_TEMPLATES.items():
    AC_ROWS.append((
        f"AC-{prefix}-001",
        f"Given an authenticated user with the required {module} permissions",
        f"When the user opens {module} and performs an authorised core action",
        "Then the page loads without exposing unauthorised data, validates the action, persists or clearly reports the result, and creates required events/audit records."
    ))
AC_ROWS.extend([
    ("AC-RBAC-001", "Given a user without financials permission", "When the user opens Dashboard, Search, Hotel Brain, Reports or a financial URL", "Then financial values and records are omitted or access is denied server-side and the attempt is handled safely."),
    ("AC-RBAC-002", "Given a disabled, pending or rejected user", "When the user attempts authentication", "Then the user receives the specific account-state message and no protected session is created."),
    ("AC-STATE-001", "Given a module query is loading, empty or fails", "When its page is displayed", "Then a distinct accessible loading, action-oriented empty or specific error state is shown."),
    ("AC-TOOLBAR-001", "Given no active call or collaboration session", "When any general module page is displayed", "Then no call/chat session toolbar appears or shifts page layout."),
    ("AC-DEMO-001", "Given data or a provider is simulated or demonstrative", "When it is displayed", "Then a visible Demo or Simulation label is shown and it is not represented as live/healthy production state."),
    ("AC-AUDIT-001", "Given a material or privileged action", "When the action succeeds or is denied", "Then the server records actor, target, action, time, result and correlation context according to policy."),
    ("AC-NOTIFY-001", "Given an event matches a notification rule", "When the event is committed", "Then the correct recipient, priority, channel and escalation rule are applied without duplicate delivery."),
])


E2E_ROWS = [
    ("1. Booking to Check-In", "Prove the stay-creation and arrival path", "Authorised front desk user; guest, room and rate data", "New guest, available/blocked rooms, booking dates", "Create booking; link guest; assign room; validate readiness/blockers; check in", "Booking checked in; room occupied; dashboard/calendar/audit updated", "Bookings, Guests, Rooms, Housekeeping, Maintenance", "booking.created/updated/checked_in; room.status_changed", "Booking and check-in audit", "Blocker/readiness notifications if applicable"),
    ("2. Check-Out to Room Turnover", "Prove departure through ready-for-sale state", "Checked-in booking and occupied room", "Booking with balance state; housekeeping staff", "Check out; mark dirty; assign clean; complete; inspect; ready", "Room available/ready and schedule/dashboard updated", "Bookings, Rooms, Housekeeping, Calendar", "booking.checked_out; room.status_changed; housekeeping.*", "Each transition", "Housekeeping assignment/delay"),
    ("3. Guest Complaint to Resolution", "Prove multi-channel service recovery", "Authorised service user and active guest", "Complaint via message/review/call/concierge", "Log; link context; classify; assign; escalate; resolve", "Closed request/task/incident with guest history and evidence", "Communications, Concierge, Reviews, Operations, Incidents", "guest.issue_logged; task/incident events", "Complaint, assignment, closure", "Owner/escalation notifications"),
    ("4. Maintenance Issue to Closure", "Prove manual/device fault lifecycle", "Maintenance users; room/device", "Priority fault and parts", "Create/receive fault; assign; block; work; complete; verify", "Fault closed and room/device restored", "Maintenance, Rooms, Smart Building, Inventory", "maintenance.*; room/device status", "Work and status changes", "Assignment/overdue/critical"),
    ("5. Security Alert to Incident Closure", "Prove restricted security response", "Security role; configured event source", "Door/access/security alert", "Receive; acknowledge; create incident; investigate; close", "Controlled closure with restricted evidence", "Security, Smart Building/CCTV, Incidents", "security.*; incident.*", "View/update/closure as policy", "Security escalation"),
    ("6. CCTV Integration to Security Display", "Prove supported camera configuration path", "Integration admin; real provider or labelled simulation", "Camera/NVR connection and area mapping", "Configure; test; import; map; open Security Center", "Camera health visible to authorised security user", "Settings, Integration Manager, CCTV, Security", "integration.*; camera.*", "Test/import/map", "Offline/test failure"),
    ("7. Smart Device to Alert Routing", "Prove device onboarding and fault routing", "Integration admin and operational responders", "Device/provider plus mapped area", "Configure; discover; map; emit event; route; acknowledge", "Correct maintenance/security/incident work is created", "Integration Manager, Smart Building, control centres", "device.*; task/incident events", "Configure/map/action", "Fault/offline/critical"),
    ("8. Inventory Low Stock to Notification", "Prove threshold and replenishment workflow", "Inventory user and department owner", "Item above threshold then adjustment below threshold", "Adjust stock; detect threshold; notify; create reorder/PO", "Low-stock state, notification and replenishment record", "Inventory, Notifications, Operations, Reports", "inventory.adjusted/low_stock/purchase_order_created", "Adjustment and PO", "Department owner/escalation"),
    ("9. Review Escalation to Operations", "Prove negative review service-recovery route", "Review and operations users", "Negative review linked to guest/stay", "Receive; classify; assign; escalate; resolve/respond", "Traceable service-recovery outcome", "Reviews, Guests, Operations, Incidents", "review.received/assigned/escalated", "Review/response/escalation", "Owner and management"),
    ("10. Enterprise Search Investigation", "Prove relevant results and permission filtering", "Users with different module permissions", "Indexed authorised and restricted records", "Search; filter; preview; open; attempt restricted query", "Only authorised results; restricted attempt handled/audited", "Enterprise Search, RBAC, source modules", "search.executed/restricted", "Query and restricted attempt per policy", "None unless policy"),
    ("11. Hotel Brain Daily Briefing", "Prove evidence-based permission-safe briefing", "Authorised GM and populated operational records", "Open blockers, incidents, guest/device/integration issues", "Generate briefing; inspect evidence; propose action; confirm/decline", "Accurate bounded briefing; controlled action; governance audit", "AI Context, Hotel Brain, all sources", "ai.answer/recommendation/decision", "Question, evidence, answer, decision", "Attention items as governed"),
    ("12. Role-Based Access Restriction", "Prove module, action and record isolation", "Representative users for every approved role", "Financial, security, audit, admin and normal records", "Navigate/direct URL/API/search/AI/export attempts", "Allowed actions succeed; denied actions reveal no data", "All modules", "access denied/restricted events", "Authentication and denials", "Security notification only if policy"),
]


def build_document() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("LaFlo | Full System Business Requirements Document")
    set_font(hr, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])

    # Editorial cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BUSINESS REQUIREMENTS DOCUMENT")
    set_font(r, size=11, bold=True, color=TEAL)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run("LaFlo Enterprise Hotel\nOperations Platform")
    set_font(r2, size=30, bold=True, color=NAVY)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(34)
    r3 = p3.add_run("Full-system business, functional, operational, security, integration and AI requirements")
    set_font(r3, size=13, color=MUTED)
    add_table(doc, ["Document", "Value"], [
        ("Version", "1.0"),
        ("Date", "26 July 2026"),
        ("Status", "Draft for product owner, QA, UAT and stakeholder review"),
        ("Prepared by", "Codex, based on the supplied platform scope and code-backed configuration audit"),
        ("Classification", "Internal - Product and Delivery"),
    ], [2200, 7160], 9.5)
    add_callout(doc, "Document purpose", "This BRD defines the intended full-platform behaviour and explicitly identifies requirements that depend on incomplete models, provider access, hardware or future implementation.", "teal")
    add_page_break(doc)

    add_heading(doc, "Contents", 1)
    for index, title in enumerate([
        "Document Control", "Executive Summary", "Business Background", "Business Objectives",
        "Current Business Problems / Gaps", "Proposed Solution", "Platform Scope", "Out of Scope",
        "User Roles and Stakeholders", "Business Process Overview", "End-to-End User Journeys",
        "Functional Requirements", "Non-Functional Requirements", "Data Requirements",
        "Integration Requirements", "Reporting and Dashboard Requirements", "AI and Automation Requirements",
        "Security and Access Control Requirements", "Audit and Logging Requirements", "Notification Requirements",
        "Event Bus Requirements", "Error Handling Requirements", "Assumptions", "Dependencies",
        "Risks and Constraints", "Acceptance Criteria", "End-to-End Testing Scope",
        "QA / UAT Considerations", "Future Enhancements", "Sign-Off Criteria"
    ], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rr = p.add_run(f"{index}. {title}")
        set_font(rr, size=10)
    add_page_break(doc)

    add_heading(doc, "1. Document Control", 1)
    add_table(doc, ["Field", "Value"], [
        ("Document name", "LaFlo Enterprise Hotel Operations Platform - Full System Business Requirements Document"),
        ("Platform name", "LaFlo Enterprise Hotel Operations Platform"),
        ("Version", "1.0"),
        ("Date", "26 July 2026"),
        ("Prepared by", "Codex / [Business Analyst name]"),
        ("Reviewed by", "[Product Owner], [Technical Lead], [QA Lead], [Security Lead]"),
        ("Approved by", "[Executive Sponsor / Steering Committee]"),
        ("Status", "Draft for review"),
    ], [2200, 7160], 9.5)
    add_heading(doc, "Change history", 2)
    add_table(doc, ["Version", "Date", "Author", "Change"], [
        ("0.1", "26 July 2026", "Codex", "Initial full-system BRD draft"),
        ("1.0", "[TBC]", "[TBC]", "Approved baseline"),
    ], [1100, 1500, 1900, 4860], 9)

    add_heading(doc, "2. Executive Summary", 1)
    add_para(doc, "LaFlo is an enterprise-grade hotel operations platform designed to centralise and automate the work that connects reservations, guest management, room readiness, housekeeping, maintenance, inventory, planning, financials, reporting, service communication, reviews, safety, security, building systems and third-party providers.")
    add_para(doc, "The platform combines a shared operational record with an Event Bus, Task Engine, Notification Engine, Audit Engine and Integration Hub. Operations Center, Security Center, Maintenance Center, Incident Center and Smart Building provide control views, while Enterprise Search, Hotel Brain and governed AI services turn authorised records into evidence-based operational insight.")
    add_callout(doc, "Business outcome", "A hotel team should be able to see what needs attention, understand the supporting evidence, act within role permissions and retain an auditable history without switching among fragmented tools.", "blue")

    add_heading(doc, "3. Business Background", 1)
    add_para(doc, "Hotel operations are time-sensitive and interdependent. A room cannot be sold or checked in safely unless reservations, cleaning, inspection and maintenance state agree. Guest complaints may arrive through calls, messages, reviews or concierge teams. Security and building events may require coordinated maintenance and incident response. Fragmented tools create delay, duplicate entry, inconsistent status and weak accountability.")
    add_para(doc, "LaFlo is intended to provide one operational system of record and one governed integration layer. The business need is not only consolidated visibility; it is reliable workflow hand-off, controlled access, evidence preservation and decision support across departments.")

    add_heading(doc, "4. Business Objectives", 1)
    add_bullets(doc, [
        "Centralise hotel operational workflows and reduce duplicate entry.",
        "Improve visibility across front desk, housekeeping, maintenance, security, finance and management.",
        "Improve room readiness and arrival/departure coordination.",
        "Improve guest experience, complaint follow-up and ownership.",
        "Reduce manual coordination through tasks, events and notifications.",
        "Improve maintenance, incident, CCTV and smart-building response.",
        "Improve inventory control and authorised financial visibility.",
        "Govern provider integrations and credential references centrally.",
        "Provide permission-filtered Enterprise Search and evidence-based Hotel Brain answers.",
        "Improve auditability, accountability and management reporting."
    ])

    add_heading(doc, "5. Current Business Problems / Gaps", 1)
    add_table(doc, ["Problem", "Business impact", "BRD response"], [
        ("Fragmented operational data", "Slow hand-offs and inconsistent status", "Shared records, event publication and linked workflows"),
        ("Manual room readiness tracking", "Arrival delays and room-sale risk", "Room, housekeeping and maintenance state validation"),
        ("Disconnected departments", "Blockers are discovered late", "Operations Center, task routing and notifications"),
        ("Limited communication visibility", "Guest issues lack ownership", "Linked messages, calls, reviews and concierge requests"),
        ("Limited security/building integration", "Faults and risks are siloed", "Security Center, Smart Building and Integration Manager"),
        ("Scattered provider credentials", "Security and support risk", "Encrypted credential references and masked UI"),
        ("Difficult cross-record investigation", "Longer resolution time", "Permission-filtered Enterprise Search"),
        ("Limited audit trail", "Weak accountability and evidence", "Server-side immutable audit requirements"),
        ("Limited AI decision support", "Managers manually aggregate risk", "Governed context, briefings and evidence-based answers"),
        ("Implementation maturity varies", "Simulation can be mistaken for live capability", "Explicit Implemented/Partial/Mocked/Planned status and labelling"),
    ], [2500, 3000, 3860], 8.7)

    add_heading(doc, "6. Proposed Solution", 1)
    add_para(doc, "LaFlo shall provide a modular web platform backed by shared hotel-scoped data, protected API routes and event-driven platform services.")
    add_table(doc, ["Layer", "Responsibility"], [
        ("Experience", "Consistent dashboard, department pages, control centres, search and AI surfaces"),
        ("Business modules", "Bookings, guests, rooms, housekeeping, inventory, calendar, financials, reports and hotel services"),
        ("Control modules", "Operations, maintenance, incidents, security and smart building"),
        ("Platform core", "Event Bus, Task Engine, Notification Engine, Audit Engine and Integration Hub"),
        ("Intelligence", "Enterprise Search, AI Context, Daily GM Briefing, Department Intelligence, Governance, Copilot and Hotel Brain"),
        ("Integration layer", "Credential-safe provider configuration, discovery, mapping, testing and health"),
        ("Data and security", "Hotel scoping, role/module/action controls, record restrictions and audit"),
    ], [2200, 7160], 9)

    add_heading(doc, "7. Platform Scope", 1)
    add_para(doc, "The following inventory defines the target business scope and records current implementation maturity. 'Partial' is not acceptance of reduced behaviour; it identifies work or dependency needed to meet the target requirement.")
    add_table(doc, ["Module", "Purpose", "Key users", "Key functions", "Records", "Links", "Current status"], MODULES,
              [1250, 1550, 1400, 1900, 1150, 1100, 1010], 7.3)

    add_heading(doc, "8. Out of Scope", 1)
    add_bullets(doc, [
        "Physical installation, cabling or certification of hotel hardware.",
        "Guaranteed support for every vendor without documented APIs, credentials and test access.",
        "Production payment or booking-channel certification unless separately contracted.",
        "Fully autonomous high-impact AI decisions without human approval.",
        "Real CCTV/NVR or smart-building validation without reachable hardware/provider access.",
        "A native mobile application or guest-facing portal unless separately approved.",
        "Advanced predictive analytics, revenue optimisation or digital-twin capability beyond the agreed release.",
        "Emergency-response replacement: LaFlo supports workflow but does not replace approved life-safety procedures."
    ])

    add_heading(doc, "9. User Roles and Stakeholders", 1)
    add_para(doc, "These are target business roles. The current application role enum is smaller; the permission model must be expanded or these roles must be represented through explicit server-owned module, action and record grants.")
    add_table(doc, ["Role", "Responsibility", "Likely modules", "Typical actions", "Restrictions"], ROLE_ROWS,
              [1250, 1850, 2450, 1950, 1860], 7.5)

    add_heading(doc, "10. Business Process Overview", 1)
    add_table(doc, ["Process", "Business outcome", "Modules"], PROCESS_ROWS, [2200, 4300, 2860], 8.5)

    add_heading(doc, "11. End-to-End User Journeys", 1)
    for title, steps, modules in JOURNEYS:
        add_heading(doc, title, 2)
        add_para(doc, f"Modules: {modules}", bold_lead="Modules:")
        add_table(doc, ["Step", "Expected business behaviour"], [(i, step) for i, step in enumerate(steps, 1)], [900, 8460], 9)

    add_heading(doc, "12. Functional Requirements", 1)
    add_para(doc, "Each requirement is independently testable. Acceptance references point to Section 26; detailed module test cases shall be derived in the subsequent QA pack.")
    fr_rows = []
    for prefix, (module, capabilities) in FR_TEMPLATES.items():
        for index, capability in enumerate(capabilities, 1):
            priority = "Must" if index <= 6 else "Should"
            roles = next((row[2] for row in MODULES if row[0].startswith(module.split(" and ")[0])), "Authorised users")
            fr_rows.append((f"{prefix}-FR-{index:03d}", f"The system shall {capability}.", priority, module, roles, f"AC-{prefix}-001"))
    add_table(doc, ["Requirement ID", "Requirement", "Priority", "Source module", "User role", "AC reference"],
              fr_rows, [1150, 3900, 700, 1200, 1500, 910], 7.5)

    add_heading(doc, "13. Non-Functional Requirements", 1)
    add_table(doc, ["ID", "Area", "Requirement", "Priority"], NFR_ROWS, [950, 1450, 6260, 700], 8.4)

    add_heading(doc, "14. Data Requirements", 1)
    add_para(doc, "Target entities are listed even when the current schema implements them through another model or does not yet persist them. Data dictionaries, retention rules and field-level classifications shall be completed during solution design.")
    add_table(doc, ["Entity", "Description", "Key fields", "Source", "Relationships", "Restrictions / gap"], DATA_ROWS,
              [1250, 1700, 2450, 1000, 1550, 1410], 7.3)
    add_heading(doc, "Data quality and lifecycle rules", 2)
    add_bullets(doc, [
        "Every operational record shall carry hotel/property scope, creation/update timestamps and a stable identifier.",
        "State-changing records shall enforce valid transitions and related-record constraints.",
        "Retention, archive, deletion and legal-hold rules shall be defined by data class.",
        "Search and AI representations shall not contain secrets and shall retain source/access references.",
        "Seed, demo and simulation records shall be segregated or visibly labelled."
    ])

    add_heading(doc, "15. Integration Requirements", 1)
    add_table(doc, ["Integration", "Purpose", "Providers", "Connection", "Data", "Triggers", "Security/error handling", "Current status"],
              INTEGRATIONS, [1050, 1300, 1250, 1050, 1100, 950, 1650, 1010], 6.9)
    add_callout(doc, "Integration state rule", "Configured, credentials detected, tested, connected, healthy, simulated and coming soon are different states and must never be collapsed into one green 'connected' status.", "amber")

    add_heading(doc, "16. Reporting and Dashboard Requirements", 1)
    add_table(doc, ["Output", "Purpose", "Audience", "Source", "Key controls"], REPORT_ROWS,
              [1550, 2600, 1700, 1750, 1760], 8)
    add_bullets(doc, [
        "Dashboard financial cards and trends shall only render and query for users with financial permission.",
        "All report filters and exports shall preserve source-module, record and action permissions.",
        "Charts shall show title, period, units, legend/tooltips where applicable and an explicit empty state.",
        "Demo, seed and simulation data shall be labelled at widget or report level.",
        "Report generation and export shall be audited with filter and format metadata."
    ])

    add_heading(doc, "17. AI and Automation Requirements", 1)
    add_table(doc, ["Capability", "Purpose", "Data sources", "Users", "Outputs", "Limitations / confirmation", "Governance"], AI_ROWS,
              [1400, 1800, 1900, 1200, 1450, 1500, 1110], 7.2)
    add_callout(doc, "AI safety boundary", "AI may explain, summarise, rank and propose. It shall not bypass source permissions, business validation or required human approval, and it shall not invent evidence.", "red")

    add_heading(doc, "18. Security and Access Control Requirements", 1)
    security_requirements = [
        ("SEC-AC-001", "The server shall be authoritative for role, module, action and record access."),
        ("SEC-AC-002", "Protected APIs shall authenticate and authorise every request; frontend visibility is not sufficient."),
        ("SEC-AC-003", "Financial, security, audit, admin and credential data shall use explicit permissions and least privilege."),
        ("SEC-AC-004", "Enterprise Search and Hotel Brain shall filter source records before preview, context or answer generation."),
        ("SEC-AC-005", "Disabled users and revoked sessions shall not access protected data."),
        ("SEC-AC-006", "Access requests shall require controlled approval/rejection and auditable password setup."),
        ("SEC-AC-007", "Credential secrets shall be encrypted or stored externally; API/UI shall expose only masked references."),
        ("SEC-AC-008", "Sensitive actions such as refunds, permission changes, integration deletion and AI execution shall require confirmation and action permission."),
        ("SEC-AC-009", "Audit-log access shall be read-only and separately restricted."),
        ("SEC-AC-010", "Hotel/property isolation shall be enforced in every query and background job."),
        ("SEC-AC-011", "Local-storage permission values shall not override an authenticated server denial."),
        ("SEC-AC-012", "Authentication shall provide specific messages for pending, rejected, disabled, password-setup and invalid-credential states."),
    ]
    add_table(doc, ["ID", "Requirement"], security_requirements, [1300, 8060], 9)

    add_heading(doc, "19. Audit and Logging Requirements", 1)
    audit_actions = [
        "Login, logout, refresh, denied access and account-state outcomes",
        "Access request submission, approval, rejection and password setup",
        "User creation, update, permission change, disable and reactivate",
        "Create, update, archive/delete and important status transitions for every business record",
        "Booking, guest, room, housekeeping, maintenance and incident changes",
        "Financial creation, status changes, refunds, adjustments and exports",
        "Settings and integration configuration, test, import, mapping and disconnect",
        "CCTV stream tests and smart-building ingestion/mapping events",
        "Enterprise Search queries and restricted attempts according to policy",
        "Hotel Brain questions, evidence references, answers and sensitive-action decisions",
        "Report generation/export and notification delivery/escalation"
    ]
    add_bullets(doc, audit_actions)
    add_para(doc, "Every audit record shall include actor/service identity, hotel, action, target type/ID, timestamp, result, correlation ID, source channel, and before/after references or change summary where safe. Audit write failure for a high-risk action shall follow an agreed fail-safe policy and generate an operational alert.")

    add_heading(doc, "20. Notification Requirements", 1)
    add_table(doc, ["Type", "Trigger", "Recipients", "Channel", "Priority", "Escalation"], NOTIFICATIONS,
              [1350, 1900, 2000, 1350, 900, 1860], 7.8)

    add_heading(doc, "21. Event Bus Requirements", 1)
    add_para(doc, "Events shall use a versioned envelope containing event ID, type, version, occurred/recorded timestamps, hotel scope, actor/service, source record, correlation/causation IDs and a data payload that excludes secrets.")
    add_table(doc, ["Category", "Example events", "Producer", "Consumers"], EVENT_ROWS,
              [1300, 3850, 1500, 2710], 8)
    add_bullets(doc, [
        "Delivery semantics, ordering needs, retry, dead-letter handling and idempotency shall be documented per event.",
        "Consumers shall not assume process-local event state.",
        "Sensitive payloads shall use references rather than raw financial, personal, security or credential data.",
        "Event publication failure shall be observable and recoverable without silently losing business state."
    ])

    add_heading(doc, "22. Error Handling Requirements", 1)
    error_rows = [
        ("Validation", "Identify field and correction; preserve safe user input"),
        ("Permission", "Return 403/not authorised without revealing record existence or content"),
        ("Missing record", "Return a specific not-found state and safe navigation"),
        ("Failed save/update", "Explain retryable vs non-retryable state and prevent duplicate commits"),
        ("Invalid/expired credentials", "Specific authentication/provider message and recovery path"),
        ("Integration failure", "Provider/category, test step, safe error code and troubleshooting action"),
        ("Camera stream/discovery", "Differentiate permission, reachability, authentication, protocol and unsupported state"),
        ("Search unavailable/no results", "Distinct unavailable and no-result states; offer safe filter/query guidance"),
        ("Hotel Brain insufficient data", "State evidence limit; do not fabricate an answer"),
        ("Restricted AI request", "Decline restricted detail and explain permission boundary without leaking content"),
        ("Event Bus/Audit/Notification failure", "Log, retry/escalate by criticality and expose operational health"),
        ("Known account state", "Pending, approved-password-not-set, rejected, disabled, invalid credentials and expired verification code use specific messages"),
    ]
    add_table(doc, ["Condition", "Required handling"], error_rows, [2400, 6960], 9)

    add_heading(doc, "23. Assumptions", 1)
    add_bullets(doc, [
        "Hotels, rooms, floors/areas, users and approved access policies exist for testing.",
        "The test environment contains controlled, non-production data covering every critical state.",
        "Event Bus, Task Engine, Notification Engine, Audit Engine and Integration Hub are available or explicitly simulated.",
        "Secure credential storage is available before production; any simulation is labelled.",
        "Real provider/hardware validation depends on network access, credentials and vendor support.",
        "AI answers are limited to current authorised records and provider availability.",
        "Demo/simulation data is clearly labelled and cannot be confused with production health.",
        "UK currency examples use GBP, e.g. £1,250.00, unless a hotel has another configured currency."
    ])

    add_heading(doc, "24. Dependencies", 1)
    add_table(doc, ["Dependency", "Why required", "Owner / evidence"], [
        ("CCTV/NVR hardware and vendor access", "Real discovery, stream and health validation", "[Integration owner] / test devices"),
        ("Smart-building devices/gateways", "Telemetry and routing validation", "[Engineering] / sandbox gateway"),
        ("Vendor APIs and credentials", "Payments, channels, weather, collaboration and AI", "[Provider owner] / sandbox credentials"),
        ("Browser camera permission", "Local-camera diagnostic test", "Tester / supported browser"),
        ("Network discovery capability", "ONVIF/device discovery", "IT / approved test network"),
        ("Secure credential storage and encryption key", "Production provider security", "Security/Platform"),
        ("Authoritative RBAC policy", "All protected journeys", "Product/Security"),
        ("Test environment and data", "Repeatable QA/UAT", "Delivery/QA"),
        ("Search indexing lifecycle", "Search/Hotel Brain freshness", "Platform team"),
        ("Observability services", "Failure diagnosis and SLO evidence", "Platform/Operations"),
    ], [2500, 4300, 2560], 8.6)

    add_heading(doc, "25. Risks and Constraints", 1)
    add_table(doc, ["Risk", "Impact", "Mitigation / decision"], [
        ("Hardware or provider unavailable", "Real E2E path cannot be proven", "Use labelled simulation for UI only; retain blocked real-integration test"),
        ("Vendor API limitations/certification", "Capability delayed or reduced", "Document provider-specific scope and certification plan"),
        ("Browser/network restrictions", "Discovery/preview fails", "Supported-browser matrix and approved network gateway"),
        ("Credential leakage", "Severe security incident", "Dedicated keys, vault/reference model, masking and audit"),
        ("Permission leakage through Search/AI/Dashboard", "Sensitive data exposure", "Server filtering, negative RBAC tests and red-team prompts"),
        ("Mock data mistaken for live", "Bad operational decisions", "Widget/provider-level Demo/Simulation labels"),
        ("Incomplete audit coverage", "Weak evidence/accountability", "Central server audit matrix and failure monitoring"),
        ("Event Bus failure/duplicates", "Lost or repeated workflow", "Outbox/idempotency/retry/dead-letter design"),
        ("Integration failure", "Degraded operations", "Isolation, health, retry and manual fallback"),
        ("Search index performance/freshness", "Slow or stale investigation", "Incremental indexing, monitoring and rebuild controls"),
        ("AI hallucination", "Unsafe guidance", "Evidence, uncertainty, human approval and governance review"),
        ("Incomplete role configuration", "Over/under-permission", "Server-owned action policy and access certification"),
        ("No live E2E integrations", "UAT confidence limited", "Separate simulated and provider-certified sign-off gates"),
    ], [2600, 3000, 3760], 8.3)

    add_heading(doc, "26. Acceptance Criteria", 1)
    add_table(doc, ["ID", "Given", "When", "Then"], AC_ROWS, [1150, 2550, 2550, 3110], 8)

    add_heading(doc, "27. End-to-End Testing Scope", 1)
    add_para(doc, "The table defines the high-level E2E scope. The next deliverable shall expand each row into positive, negative, boundary, permission, integration-degraded and recovery test cases with traceability to functional requirements.")
    add_table(doc, ["Area", "Objective", "Preconditions", "Data", "High-level steps", "Expected outcome", "Modules", "Events", "Audit", "Notifications"],
              E2E_ROWS, [900, 1050, 1100, 900, 1300, 1150, 950, 800, 650, 560], 6.4)

    add_heading(doc, "28. QA / UAT Considerations", 1)
    qa_items = [
        ("Navigation and routes", "Every permitted module is reachable; hidden/direct routes remain protected"),
        ("Forms and validation", "Required, boundary, invalid, duplicate and preservation behaviour"),
        ("Status transitions", "Valid and invalid state changes with event/audit verification"),
        ("Dashboards and charts", "Permission, data provenance, empty/error/loading and drill-through"),
        ("Tables and filters", "Search, applied filters, clear, pagination, responsive scroll and row actions"),
        ("Enterprise Search", "Relevance, freshness, category filters and negative permission cases"),
        ("Hotel Brain / AI", "Evidence, uncertainty, restricted prompts, confirmation and governance logs"),
        ("Event Bus", "Schema, correlation, idempotency, retry and dead-letter cases"),
        ("Notifications", "Recipient, channel, priority, duplicates, acknowledgement and escalation"),
        ("Audit", "Coverage, immutability, access and failure handling"),
        ("RBAC", "Module/action/record matrix across all approved roles and API/direct URL/export attempts"),
        ("Integrations", "Real vs simulation, credential masking, test/import/map/health/disconnect and failure isolation"),
        ("UX states", "Loading, empty, error, permission-denied, success and specific known states"),
        ("Responsive/accessibility", "Desktop/tablet/smaller laptop; keyboard, focus, labels and contrast"),
        ("Compatibility", "Supported browser versions and provider/browser permissions"),
        ("Build/runtime", "Frontend/backend builds, startup, health, console and API errors"),
    ]
    add_table(doc, ["Checklist area", "Coverage"], qa_items, [2500, 6860], 9)
    add_callout(doc, "UAT evidence", "A simulated integration can approve interface behaviour but cannot satisfy a real provider-connectivity acceptance criterion. Evidence must identify environment, data source and provider state.", "amber")

    add_heading(doc, "29. Future Enhancements", 1)
    add_bullets(doc, [
        "Native mobile app and guest-facing portal.",
        "Multi-property enterprise management and consolidated reporting.",
        "Advanced revenue management, forecasting and demand optimisation.",
        "Predictive maintenance and building digital twin.",
        "Additional certified CCTV, lock, BMS, payment and booking-channel providers.",
        "Advanced workflow orchestration and configurable SLA/routing policy.",
        "Advanced analytics, compliance reporting and evidence retention.",
        "Durable Hotel Brain conversations, evaluation dashboards and model/provider routing."
    ])

    add_heading(doc, "30. Sign-Off Criteria", 1)
    signoff = [
        "BRD reviewed by Product Owner and module owners.",
        "Functional and non-functional requirements approved.",
        "Data classification, retention and access model approved.",
        "Security and AI governance requirements approved.",
        "Integration statuses, dependencies and assumptions accepted.",
        "Out-of-scope items and future roadmap accepted.",
        "Full E2E QA test pack created and traceable to requirement IDs.",
        "Critical test cases passed in the agreed environment.",
        "High-severity defects resolved or formally accepted with owner/date.",
        "Real versus simulated integration evidence identified.",
        "Stakeholder sign-off recorded."
    ]
    add_table(doc, ["Sign-off item", "Owner", "Status / evidence"], [(item, "[TBC]", "[ ] Pending") for item in signoff],
              [5700, 1600, 2060], 9)

    add_heading(doc, "Completion Summary", 1)
    add_para(doc, "This BRD defines the intended business behaviour for the full LaFlo platform, with module scope, roles, processes, detailed functional and non-functional requirements, target data entities, provider integration status, AI governance, security, audit, notifications, events, acceptance criteria and twelve E2E test areas.")
    add_heading(doc, "Key assumptions", 2)
    add_para(doc, "Controlled test data, server-authoritative permissions, platform engines and secure credential storage are available; live integration acceptance depends on provider/hardware access.")
    add_heading(doc, "Key risks", 2)
    add_para(doc, "The most material risks are permission leakage through dashboard/search/AI, mock data being mistaken for live state, incomplete action-level RBAC, integration maturity variation, missing call/report/AI persistence, and lack of live hardware certification.")
    add_heading(doc, "Recommended next step", 2)
    add_callout(doc, "Create the Full End-to-End QA Test Pack", "Translate every functional requirement and acceptance criterion into traceable positive, negative, permission, boundary, integration-degraded, event, audit and notification test cases.", "teal")

    doc.core_properties.title = "LaFlo Enterprise Hotel Operations Platform - Full System BRD"
    doc.core_properties.subject = "Full-system business requirements for QA, UAT, handover and sign-off"
    doc.core_properties.author = "LaFlo Product and Delivery"
    doc.core_properties.keywords = "LaFlo, hotel operations, BRD, QA, UAT, E2E, AI, integrations"
    doc.core_properties.comments = "Generated from the approved platform scope and code-backed configuration audit."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
