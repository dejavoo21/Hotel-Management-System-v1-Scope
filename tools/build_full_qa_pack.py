from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_full_system_brd as brd


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "LaFlo_Full_End_to_End_QA_Test_Pack_v1.0.docx"


MODULE_CASES = {
    "DASH": ("Dashboard", [
        "Dashboard loads successfully", "KPI cards display correctly", "Today's arrivals display",
        "Today's departures display", "Occupancy displays", "ADR displays where authorised",
        "Revenue summary is restricted", "Room readiness summary displays", "Rooms not ready display",
        "Housekeeping summary displays", "Maintenance issues display", "Open incidents display",
        "Guest experience alerts display", "Security and CCTV alerts display", "Smart Building health displays",
        "Integration health displays", "Tasks panel displays", "Recent activities display",
        "Enterprise Search shortcut works", "Hotel Brain shortcut works", "Role-based visibility works",
        "Demo and simulation data is labelled",
    ]),
    "BOOK": ("Bookings / Reservations", [
        "Create booking", "View booking", "Update booking", "Cancel booking", "Assign room",
        "Link guest to booking", "Add special request", "Check in guest", "Check out guest",
        "Booking appears on calendar", "Booking updates dashboard", "Booking links to room and guest",
        "Booking status creates audit log", "Payment status is permission controlled",
    ]),
    "GUEST": ("Guests", [
        "Create guest profile", "Update guest profile", "View guest profile", "Link guest to booking",
        "Add guest preference", "Add guest complaint", "View stay history", "View linked messages",
        "View linked calls", "View linked reviews", "View linked concierge requests",
        "Privacy permissions are respected", "Guest update creates audit log",
    ]),
    "ROOM": ("Rooms", [
        "View room list", "View room details", "Update occupancy status", "Mark room clean",
        "Mark room dirty", "Mark pending inspection", "Mark room ready", "Mark room out of service",
        "Link room to booking", "Link room to guest", "Link housekeeping task", "Link maintenance task",
        "View linked smart devices", "Room status updates dashboard",
    ]),
    "HK": ("Housekeeping", [
        "Create housekeeping task", "Assign housekeeping staff", "Update cleaning status",
        "Mark room cleaned", "Mark room inspected", "Escalate blocker", "Link task to room",
        "Link task to booking", "Link task to maintenance issue", "Update affects room readiness",
        "Delay triggers notification", "Update creates audit log",
    ]),
    "INV": ("Inventory", [
        "Create inventory item", "Update inventory item", "Adjust stock level", "Record stock movement",
        "Set reorder threshold", "Trigger low-stock alert", "Create reorder task or suggestion",
        "Link inventory to department", "View supplier information", "Update creates audit log",
    ]),
    "CAL": ("Calendar", [
        "View reservation calendar", "View check-ins", "View check-outs", "View housekeeping schedule",
        "View maintenance schedule", "Create calendar event", "Update calendar event",
        "Delete calendar event where authorised", "Calendar reflects booking changes",
        "Calendar reflects task scheduling", "Calendar links to source records",
    ]),
    "FIN": ("Financials", [
        "Authorised user views financial summary", "Unauthorised user cannot view financial data",
        "View invoice or payment record", "Link payment to booking", "Link payment to guest",
        "Update payment status", "View outstanding payments", "View refunds where applicable",
        "Export financial data", "Financial update creates audit log",
        "Search hides financial records from unauthorised user", "Hotel Brain hides financial data",
    ]),
    "REP": ("Reports", [
        "Reports page loads", "Report catalogue displays", "Filter reports by category",
        "View report details", "Generate report", "Preview report", "Export CSV", "Export PDF",
        "Finance reports are restricted", "Security reports are restricted", "Audit reports are restricted",
        "Generation and export create audit log",
    ]),
    "REV": ("Reviews", [
        "Create or import review", "View review", "Link review to guest and booking",
        "Assign review response", "Update response status", "Mark review responded",
        "Escalate negative review to Operations", "Escalate review to Incident Center",
        "Review appears on guest experience dashboard", "Update creates audit log",
    ]),
    "CON": ("Concierge", [
        "Create guest request", "Assign request", "Update request status", "Add request note",
        "Escalate to Operations", "Escalate to Incident Center", "Link request to guest",
        "Link request to room", "Link request to booking", "Create linked housekeeping or maintenance task",
        "Close request", "Update creates audit log",
    ]),
    "MSG": ("Messages", [
        "Messages page loads", "Thread list displays", "Conversation panel displays", "Create message record",
        "Reply to message", "Add internal note", "Link message to guest", "Link message to booking",
        "Link message to room", "Create follow-up action", "Escalate to concierge request",
        "Escalate to incident", "Chat toolbar appears only in active session",
        "Chat toolbar does not appear globally",
    ]),
    "CALL": ("Calls", [
        "Calls page loads", "Call log displays", "Create call log", "Add call notes",
        "Link call to guest", "Link call to booking", "Link call to room", "Create follow-up task",
        "Escalate call to incident", "Start active call session", "Call toolbar appears only in active session",
        "Call toolbar does not appear globally", "Call update creates audit log",
    ]),
    "USER": ("User Management", [
        "Create user", "Update user", "Disable user", "Reactivate user", "Assign role",
        "Update department", "Update permissions", "View pending access requests",
        "Approve access request", "Reject access request", "Approved user enters password setup",
        "Pending password setup status displays", "Active user signs in",
        "Disabled and rejected users cannot sign in", "Management action creates audit log",
    ]),
    "OPS": ("Operations Center", [
        "Operations Center loads", "Cross-department alerts display", "Operational blockers display",
        "Guest experience issues display", "Daily operational status displays", "Assign operational task",
        "Escalate blocker", "Close operational alert", "Source modules update Operations Center",
    ]),
    "MAINT": ("Maintenance Center", [
        "Create maintenance task", "Assign technician", "Update maintenance status",
        "Link task to room", "Link task to device", "Link task to incident", "Mark task resolved",
        "Close task", "Device fault creates or suggests task", "Overdue work displays on dashboard",
        "Update creates audit log",
    ]),
    "INC": ("Incident Center", [
        "Create incident", "Classify incident", "Set severity", "Set priority", "Assign owner",
        "Link incident to guest", "Link incident to room", "Link incident to booking",
        "Link incident to device or camera", "Add investigation note", "Add evidence",
        "Escalate incident", "Close incident", "Incident appears on dashboard",
        "Update creates audit log",
    ]),
    "SEC": ("Security Center", [
        "Security Center loads", "CCTV overview displays", "Camera list or grid displays",
        "Active cameras display", "Offline cameras display", "Security alerts display",
        "Access and door events display", "Restricted-area alerts display",
        "Escalate security alert to incident", "CCTV integrations display",
        "Smart lock and sensor events display", "Raw credentials are not exposed",
    ]),
    "SB": ("Smart Building", [
        "Smart Building page loads", "Device inventory displays", "Device health displays",
        "Sensors display", "Doors and access view displays", "Energy view displays", "HVAC view displays",
        "Assets display", "Building health score displays", "Map device to room floor or area",
        "Offline status displays", "Low-battery status displays", "Fault routes to control module",
        "Raw credentials are not exposed",
    ]),
    "INT": ("Integration Manager", [
        "Open Settings Integrations", "Integration categories display", "Provider cards display",
        "Configure integration", "Test connection", "View health status", "View logs",
        "Disable integration", "Credentials are masked", "Raw credentials are not displayed",
        "Integration action creates audit log", "Demo and simulation providers are labelled",
    ]),
}


CCTV_CASES = [
    "Manual RTSP camera setup", "HLS camera setup", "MJPEG camera setup", "Snapshot URL setup",
    "ONVIF manual camera setup", "Local USB browser preview", "IP camera discovery",
    "NVR connection", "NVR channel import", "Map camera to hotel area",
    "Camera appears in Security Center", "Failed stream test shows clear error",
    "Simulated discovery is clearly labelled",
]
SMART_INT_CASES = [
    "Smart lock provider setup", "Sensor provider setup", "HVAC provider setup",
    "Energy meter provider setup", "Access control provider setup", "Manual device creation",
    "Local network discovery", "Gateway connection", "Webhook setup", "Device import",
    "Device mapping", "Device health status", "Device alert routing", "Simulated hardware is labelled",
]
SEARCH_CASES = [
    "Global search", "Advanced search", "Search by guest", "Search by room", "Search by booking",
    "Search by incident", "Search by maintenance task", "Search by CCTV camera",
    "Search by smart building device", "Search by message", "Search by call", "Search by review",
    "Search financial record where authorised", "Search user where authorised", "Search integration",
    "Search audit log where authorised", "Category grouping", "Filters", "Result preview",
    "Saved searches", "Recent searches", "Permission filtering", "No-results state", "Error state",
]
AI_CASES = [
    "Ask for daily GM briefing", "Ask what happened overnight", "Ask for rooms not ready",
    "Ask for unresolved incidents", "Ask for offline devices", "Ask for overdue maintenance",
    "Ask for guest complaints in last 7 days", "Ask for CCTV health issues",
    "Ask for smart building risks", "Unauthorised financial or security request is refused",
    "Answer includes evidence and source records", "Insufficient-data response is bounded",
    "Suggested action is shown", "Sensitive action requires confirmation",
    "AI answer creates audit log", "AI governance metadata is recorded",
]
NEG_CASES = [
    "Required field missing", "Invalid date", "Invalid room selection", "Duplicate guest",
    "Invalid booking status change", "Unauthorised access", "Failed save", "Failed update",
    "Failed integration connection", "Invalid credentials", "Expired credentials",
    "Failed camera stream", "No cameras discovered", "Device discovery failed",
    "Search unavailable", "Hotel Brain unavailable", "Hotel Brain restricted request",
    "Event Bus failure", "Audit log failure", "Notification failure",
    "Known authentication account states use specific messages",
]
REG_CASES = [
    "Login and access", "Request access", "Password setup", "Dashboard load", "Booking creation",
    "Guest creation", "Room status update", "Housekeeping status update",
    "Maintenance task creation", "Incident creation", "Inventory low-stock",
    "Integration Manager load", "CCTV integration setup", "Smart Building device display",
    "Enterprise Search query", "Hotel Brain query", "Audit log creation",
    "Notification creation", "Role-based restriction",
]
UAT_ROLES = [
    "General Manager", "Front Desk Staff", "Housekeeping Manager", "Maintenance Manager",
    "Security Manager", "IT / Integration Administrator", "Finance User",
    "Operations Manager", "Concierge Staff",
]


def tc(case_id, area, title, req, priority, test_type, preconditions, data, steps, expected):
    return [
        case_id, area, title, req, priority, test_type, preconditions, data, steps, expected,
        "", "Not Run", "", "",
    ]


def standard_case(prefix, index, area, title, test_type="Functional", priority="High"):
    action = title[0].lower() + title[1:]
    return tc(
        f"{prefix}-TC-{index:03d}", area, title, f"{prefix}-FR-{min(index, 7):03d}",
        priority, test_type,
        f"An active test user has the required {area} module and action permissions; required source records exist.",
        f"Controlled {area} records covering normal, empty and restricted states.",
        f"1. Sign in as the authorised test user.\n2. Open {area}.\n3. {title}.\n4. Save or confirm the action where applicable.\n5. Refresh and inspect linked records, events, notifications and audit evidence.",
        f"The system shall {action}; persist and display the correct state; preserve permissions; show a specific error if unsupported; and create required audit/event records.",
    )


def build_cases():
    functional = []
    for prefix, (area, titles) in MODULE_CASES.items():
        for i, title in enumerate(titles, 1):
            functional.append(standard_case(prefix, i, area, title))

    cctv = [standard_case("CCTV", i, "CCTV Integrations", title, "Integration", "Critical" if "credential" in title.lower() else "High") for i, title in enumerate(CCTV_CASES, 1)]
    smart = [standard_case("SBINT", i, "Smart Building Integrations", title, "Integration") for i, title in enumerate(SMART_INT_CASES, 1)]
    integrations = cctv + smart
    for i, row in enumerate(brd.INTEGRATIONS, len(integrations) + 1):
        integrations.append(standard_case("INT", i, "Integration Manager", f"{row[0]} provider status and failure handling", "Integration"))

    e2e = []
    for i, row in enumerate(brd.E2E_ROWS, 1):
        e2e.append(tc(
            f"E2E-TC-{i:03d}", row[0], row[1], "Multiple BRD functional requirements",
            "Critical", "E2E", row[2], row[3],
            f"1. Prepare the stated preconditions and data.\n2. Execute: {row[4]}.\n3. Validate source and linked records.\n4. Validate expected events: {row[7]}.\n5. Validate audit: {row[8]}.\n6. Validate notifications: {row[9]}.",
            row[5],
        ))

    rbac = []
    roles = [r[0] for r in brd.ROLE_ROWS if r[0] not in {"Department Manager"}]
    for i, role in enumerate(roles, 1):
        rbac.append(tc(
            f"RBAC-TC-{i:03d}", "Role-Based Access", f"Validate {role} access matrix",
            "SEC-AC-001 to SEC-AC-012", "Critical", "RBAC",
            f"An active {role} account and a comparison System Administrator account exist.",
            "Authorised and restricted module, financial, security, audit, user, integration and AI records.",
            "1. Sign in as the role.\n2. Verify sidebar and direct-route module visibility.\n3. Test create, read, update and delete/archive actions.\n4. Search authorised and restricted records.\n5. Ask Hotel Brain authorised and restricted questions.\n6. Verify Dashboard, audit, financial, CCTV, Settings and Integration visibility.\n7. Repeat restricted attempts through API/export where applicable.",
            "Only the approved module, action and record scope is available. Restricted data is absent from UI, direct routes, APIs, exports, Search and Hotel Brain; required denials are audited.",
        ))

    events = []
    for i, row in enumerate(brd.EVENT_ROWS, 1):
        events.append(tc(
            f"EVENT-TC-{i:03d}", "Event Bus", f"Publish and consume {row[0]} events",
            "Section 21 Event Bus Requirements", "High", "Integration",
            "Event capture/inspection is enabled and the source action can be executed.",
            f"Source records for {row[0]} and a correlation identifier.",
            f"1. Execute each applicable source transition: {row[1]}.\n2. Inspect the event envelope and payload.\n3. Verify hotel scope, version, timestamps, actor, source, correlation and causation.\n4. Verify consumers: {row[3]}.\n5. Replay or retry to confirm idempotency where applicable.",
            "Exactly the required versioned events are published and consumed without secrets, duplication or loss; failures are observable and recoverable.",
        ))

    audit_titles = [
        "Login and access outcome", "Access request approval and rejection", "User create update and disable",
        "Record creation", "Record update", "Record delete or archive", "Booking status change",
        "Guest profile change", "Room status change", "Housekeeping update", "Maintenance update",
        "Incident update", "Financial update", "Settings change", "Integration setup",
        "Connection test", "Device import", "CCTV stream test", "Search query",
        "Restricted access attempt", "Hotel Brain query", "Hotel Brain answer",
        "Sensitive action confirmation and rejection", "Report generation and export",
    ]
    audit = [standard_case("AUDIT", i, "Audit Engine", title, "Functional", "Critical" if any(k in title.lower() for k in ["financial", "restricted", "user", "sensitive"]) else "High") for i, title in enumerate(audit_titles, 1)]
    notifications = [standard_case("NOTIF", i, "Notification Engine", row[0], "Functional", "Critical" if row[4] == "Critical" else "High") for i, row in enumerate(brd.NOTIFICATIONS, 1)]
    search = [standard_case("SEARCH", i, "Enterprise Search", title, "Functional", "Critical" if "permission" in title.lower() or "financial" in title.lower() else "High") for i, title in enumerate(SEARCH_CASES, 1)]
    ai = [standard_case("AI", i, "Hotel Brain and AI Governance", title, "Functional", "Critical" if "unauthorised" in title.lower() or "confirmation" in title.lower() else "High") for i, title in enumerate(AI_CASES, 1)]
    negative = [standard_case("NEG", i, "Cross-Platform Negative Testing", title, "Negative", "Critical" if any(k in title.lower() for k in ["unauthorised", "credential", "audit", "event bus"]) else "High") for i, title in enumerate(NEG_CASES, 1)]
    regression = [standard_case("REG", i, "Critical Regression", title, "Regression", "Critical") for i, title in enumerate(REG_CASES, 1)]
    uat = []
    for i, role in enumerate(UAT_ROLES, 1):
        uat.append(tc(
            f"UAT-TC-{i:03d}", "Business UAT", f"{role} completes a representative working-day scenario",
            "BRD business processes and role requirements", "High", "UAT",
            f"A trained {role} tester and realistic department data are available.",
            "Representative records, alerts, tasks and linked evidence for the role.",
            f"1. Sign in as {role}.\n2. Review the role dashboard and attention items.\n3. Complete the role's main daily workflow.\n4. Follow a linked record into another authorised module.\n5. Resolve or escalate an exception.\n6. Confirm notifications and activity/audit evidence.\n7. Confirm restricted information is not shown.",
            f"The {role} can complete the business scenario efficiently with correct data, ownership, hand-off, permissions and evidence.",
        ))
    return {
        "Functional Cases": functional, "E2E Scenarios": e2e, "RBAC Cases": rbac,
        "Integration Cases": integrations, "Event Bus Cases": events, "Audit Cases": audit,
        "Notification Cases": notifications, "Search Cases": search, "AI Governance Cases": ai,
        "Negative Cases": negative, "Regression Suite": regression, "UAT Suite": uat,
    }


HEADERS = [
    "Test Case ID", "Test Area / Module", "Test Title", "Requirement Reference",
    "Priority", "Test Type", "Preconditions", "Test Data", "Test Steps", "Expected Result",
    "Actual Result", "Status", "Defect ID", "Comments",
]
WIDTHS = [650, 800, 1100, 900, 450, 550, 1100, 900, 2100, 1600, 1000, 550, 550, 750]


def configure_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)


def add_case_table(doc, cases):
    brd.add_table(doc, HEADERS, cases, WIDTHS, 6.5)


def build():
    cases = build_cases()
    doc = Document()
    brd.configure_styles(doc)
    configure_landscape(doc)
    hp = doc.sections[0].header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    brd.set_font(hp.add_run("LaFlo | Full End-to-End QA Test Pack"), size=8, bold=True, color=brd.MUTED)
    brd.add_page_number(doc.sections[0].footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brd.set_font(p.add_run("FULL END-TO-END QA TEST PACK"), size=12, bold=True, color=brd.TEAL)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brd.set_font(p2.add_run("LaFlo Enterprise Hotel Operations Platform"), size=28, bold=True, color=brd.NAVY)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brd.set_font(p3.add_run("Functional, E2E, regression, UAT, RBAC, integration, AI, event, audit and notification validation"), size=12, color=brd.MUTED)
    total = sum(len(v) for v in cases.values())
    brd.add_callout(doc, "Execution baseline", f"Version 1.0 | 26 July 2026 | {total} executable test cases | Initial status: Not Run", "teal")
    brd.add_page_break(doc)

    brd.add_heading(doc, "1. QA Test Pack Overview", 1)
    brd.add_para(doc, "This pack translates the Full System BRD into executable manual test cases for functional, end-to-end, regression, UAT, RBAC, provider workflow, AI governance, Enterprise Search, audit, notification and stakeholder sign-off.")
    brd.add_table(doc, ["Control", "Value"], [
        ("Platform", "LaFlo Enterprise Hotel Operations Platform"),
        ("QA Pack", "Full End-to-End QA Test Pack"),
        ("Version / date", "1.0 / 26 July 2026"),
        ("BRD baseline", "LaFlo Full System BRD v1.0"),
        ("Initial execution state", "All test cases Not Run"),
        ("Prepared by", "Codex / [QA Lead]"),
        ("Reviewers", "[Product Owner], [Engineering Lead], [Security Lead], [QA Lead]"),
    ], [2200, 12200], 9)

    brd.add_heading(doc, "2. Test Scope", 1)
    brd.add_para(doc, "Scope includes the platform core, all business and control modules, intelligence and AI capabilities, provider integration workflows, role restrictions, cross-module journeys, UI states, event publication, notifications, audit evidence and key non-functional controls.")
    brd.add_heading(doc, "3. Out of Scope", 1)
    brd.add_bullets(doc, [
        "Physical installation or certification of unavailable hotel hardware.",
        "Production payment/channel certification without provider sandboxes and credentials.",
        "Native mobile or guest portal testing unless separately delivered.",
        "Unimplemented planned-provider behaviour beyond interface and Coming Soon validation.",
    ])
    brd.add_heading(doc, "4. Test Assumptions", 1)
    brd.add_bullets(doc, [
        "The BRD is the approved requirements baseline and code/configuration gaps remain visible as blocked or expected-fail tests.",
        "Representative hotel, user, room, guest, booking, financial and operational test data exists.",
        "Demo/simulation data and providers are clearly labelled.",
        "All tests use non-production data and controlled credentials.",
    ])
    brd.add_heading(doc, "5. Test Dependencies", 1)
    brd.add_table(doc, ["Dependency", "Required evidence"], [
        ("Frontend/API builds", "Successful build and startup logs"),
        ("Database and seed data", "Repeatable reset/seed procedure"),
        ("RBAC accounts", "Named role/permission matrix and credentials"),
        ("Event/audit/notification inspection", "Safe access to queues/logs/records"),
        ("Provider sandboxes/hardware", "Credentials, network access and test resources"),
        ("Search/AI providers", "Indexed records, provider state and governance access"),
    ], [3800, 10600], 9)
    brd.add_heading(doc, "6. Test Environment Requirements", 1)
    brd.add_table(doc, ["Environment item", "Requirement"], [
        ("QA URL/API", "Stable isolated deployment with health endpoints"),
        ("Browsers", "Current Chrome and Edge; Safari where agreed"),
        ("Data reset", "Documented rollback/reseed or isolated test records"),
        ("Observability", "Correlation IDs, API logs, event, audit and notification inspection"),
        ("Integrations", "Separate real, sandbox, simulated and unavailable statuses"),
        ("Security", "No production secrets; credential masking enabled"),
    ], [3800, 10600], 9)
    brd.add_heading(doc, "7. Test Data Requirements", 1)
    brd.add_para(doc, "Prepare normal, boundary, empty, invalid, duplicate, overdue, blocked, failed, restricted, simulated and unavailable records for every relevant module. Test data must include GBP examples, multiple room/booking states, restricted financial/security/audit records, indexed and non-indexed records, device/camera health states and AI evidence/no-evidence scenarios.")
    brd.add_heading(doc, "8. User Roles Required for Testing", 1)
    brd.add_table(doc, ["Role", "Purpose"], [(r[0], r[1]) for r in brd.ROLE_ROWS if r[0] != "Department Manager"], [3800, 10600], 8.5)
    brd.add_heading(doc, "9. Test Execution Guidelines", 1)
    brd.add_bullets(doc, [
        "Record Actual Result, Status, Defect ID and Comments for every executed case.",
        "Capture screenshots, API response, correlation/event ID and audit/notification evidence when applicable.",
        "Do not mark a simulated integration as passing a real-provider acceptance criterion.",
        "Re-run failed cases after fixes and execute the linked regression cases.",
        "Stop and escalate any suspected personal, financial, security or credential exposure.",
    ])
    brd.add_heading(doc, "10. Defect Severity and Priority Definitions", 1)
    brd.add_table(doc, ["Level", "Severity definition", "Priority guidance"], [
        ("Critical", "Security/data exposure, data loss, system unavailable or critical hotel workflow blocked", "P0 immediate triage"),
        ("High", "Major module/E2E failure with no acceptable workaround", "P1 current cycle"),
        ("Medium", "Material issue with workaround or limited scope", "P2 planned fix"),
        ("Low", "Cosmetic, copy or low-impact usability issue", "P3 backlog"),
    ], [1800, 6900, 5700], 8.5)
    brd.add_heading(doc, "11. Test Case Naming Convention", 1)
    brd.add_para(doc, "IDs use AREA-TC-NNN. Functional prefixes follow the BRD (DASH, BOOK, GUEST, ROOM, HK, INV, CAL, FIN, REP, REV, CON, MSG, CALL, USER, OPS, MAINT, INC, SEC, SB, INT). Cross-cutting prefixes include CCTV, E2E, RBAC, EVENT, AUDIT, NOTIF, SEARCH, AI, NEG, REG and UAT.")

    section_map = {
        "Functional Cases": "12. Functional Test Cases by Module",
        "E2E Scenarios": "13. End-to-End Test Scenarios",
        "RBAC Cases": "14. Role-Based Access Control Test Cases",
        "Integration Cases": "15. Integration Test Cases",
        "Event Bus Cases": "16. Event Bus Test Cases",
        "Audit Cases": "17. Audit Logging Test Cases",
        "Notification Cases": "18. Notification Test Cases",
        "Search Cases": "19. Enterprise Search Test Cases",
        "AI Governance Cases": "20. Hotel Brain and AI Governance Test Cases",
        "Negative Cases": "21. Error Handling and Negative Test Cases",
        "Regression Suite": "22. Regression Test Suite",
        "UAT Suite": "23. UAT Test Suite",
    }
    for key, title in section_map.items():
        brd.add_page_break(doc)
        brd.add_heading(doc, title, 1)
        brd.add_para(doc, f"{len(cases[key])} cases. Execution columns are intentionally blank except Status = Not Run.")
        add_case_table(doc, cases[key])

    brd.add_page_break(doc)
    brd.add_heading(doc, "24. Test Execution Summary Template", 1)
    brd.add_table(doc, ["Field", "Value / Formula"], [
        ("Test cycle", "[Cycle name]"), ("Date", "[YYYY-MM-DD]"), ("Tester", "[Name]"),
        ("Environment", "[QA/UAT/Browser/build]"), ("Total test cases", "[Count]"),
        ("Passed", "[Count]"), ("Failed", "[Count]"), ("Blocked", "[Count]"),
        ("Not run", "[Count]"), ("Pass percentage", "[Passed / executed]"),
        ("Open critical defects", "[Count]"), ("Open high defects", "[Count]"),
        ("Key risks", "[Summary]"), ("Recommendation", "[Proceed / Conditional / Stop]"),
        ("Sign-off status", "[Pending / Approved / Rejected]"),
    ], [4000, 10400], 9)

    brd.add_heading(doc, "25. Defect Log Template", 1)
    defect_headers = ["Defect ID", "Date", "Raised by", "TC ID", "Module", "Title", "Description", "Steps", "Expected", "Actual", "Severity", "Priority", "Evidence", "Assigned to", "Status", "Retest date", "Retest result", "Closure comments"]
    brd.add_table(doc, defect_headers, [["" for _ in defect_headers] for _ in range(8)],
                  [550, 600, 650, 600, 650, 900, 1200, 1300, 950, 950, 600, 600, 700, 700, 650, 650, 700, 900], 6.2)

    brd.add_heading(doc, "26. Sign-Off Checklist", 1)
    checklist = [
        "QA pack reviewed and requirement traceability accepted",
        "Environment, data and role accounts approved",
        "Critical E2E, RBAC, security and data-leakage tests passed",
        "Regression suite passed",
        "Provider tests distinguish real, sandbox and simulation evidence",
        "Open Critical defects = 0",
        "Open High defects resolved or formally accepted",
        "UAT completed by nominated business roles",
        "Risks, blocked tests and workarounds accepted",
        "Product Owner, QA Lead, Engineering Lead and Security Lead sign-off recorded",
    ]
    brd.add_table(doc, ["Checklist item", "Owner", "Status", "Evidence / comments"],
                  [[item, "[TBC]", "[ ] Pending", ""] for item in checklist], [6500, 1800, 1700, 4400], 8.5)
    brd.add_heading(doc, "Completion Summary", 1)
    brd.add_para(doc, f"This pack contains {total} executable test cases across functional, E2E, RBAC, integrations, Event Bus, audit, notifications, Enterprise Search, Hotel Brain/AI governance, negative, regression and UAT coverage.")
    brd.add_heading(doc, "Recommended execution order", 2)
    brd.add_para(doc, "1) Build/health/authentication; 2) RBAC and sensitive-data controls; 3) core booking-room-housekeeping journeys; 4) control centres; 5) integrations and events; 6) Search and AI; 7) regression; 8) UAT.")
    brd.add_heading(doc, "Defect triage recommendation", 2)
    brd.add_para(doc, "Triage Critical and High defects daily with Product, Engineering, QA and Security. Any credential, financial, guest, security, Search or AI permission leakage is a release blocker.")

    doc.core_properties.title = "LaFlo Full End-to-End QA Test Pack"
    doc.core_properties.subject = "Execution-ready QA, E2E, regression, RBAC, integration and UAT test pack"
    doc.core_properties.author = "LaFlo Product and Delivery"
    doc.save(OUT)
    print(OUT)
    print({k: len(v) for k, v in cases.items()})
    print("total", total)


if __name__ == "__main__":
    build()
